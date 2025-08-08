#!/usr/bin/env python3

from universal_md_converter import UniversalMDConverter
import os

def generate_initial_version_check():
    """초기 버전으로 들여쓰기 확인용 문서 생성"""
    print('🔄 초기 버전으로 들여쓰기 확인용 새 파일명으로 생성...')
    
    converter = UniversalMDConverter()
    
    # 기본 convert 메소드를 사용하되, 출력 파일명만 변경
    md_file = 'RWSL_항공시스템_사업계획서.md'
    
    print(f'🔄 변환 시작: {md_file}')
    
    with open(md_file, 'r', encoding='utf-8') as f:
        content = f.read()
        
    lines = content.split('\n')
    
    # 기존 convert 로직과 동일하게 처리
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    
    # 제목 먼저 추가
    title_para = converter.document.add_paragraph('실화상 카메라 기반 장거리 드론탐지시스템 구축 사업계획서')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    
    converter.document.add_page_break()
    
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        
        if not line:  # 빈 줄
            i += 1
            continue
            
        if line.startswith('# '):  # H1 제목 (문서 제목이므로 스킵)
            i += 1
            continue
            
        elif line.startswith('## '):  # H2 제목  
            title = line[3:].strip()
            if title == '주석':  # 주석 섹션은 따로 처리
                i = converter.process_footnote_section(lines, i)
                continue
            else:
                para = converter.document.add_paragraph(title)
                para.style = 'CustomHeading1'
                
        elif line.startswith('### '):  # H3 제목
            title = line[4:].strip()
            para = converter.document.add_paragraph(title)
            para.style = 'CustomHeading2'
            
        elif line.startswith('!['):  # 이미지
            i = converter.process_image(lines, i)
            continue
            
        elif line.startswith('<그림') or line.startswith('<표'):  # 캡션
            para = converter.document.add_paragraph(line)
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.bold = True
            
        elif line.startswith('|') and '|' in line:  # 테이블
            i = converter.process_table(lines, i)
            continue
            
        elif line.startswith('□') or line.startswith('○'):  # 불릿 포인트
            converter.add_bullet_paragraph(line)
            
        elif line.strip() != '---':  # 일반 텍스트 (구분선 제외)
            if line.strip():  # 빈 줄이 아닌 경우만
                para = converter.document.add_paragraph(line)
                run = para.runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(11)
                
        i += 1
    
    # 확인용 새 파일명으로 저장
    output_file = 'output/RWSL_항공시스템_사업계획서_초기버전_들여쓰기확인.docx'
    os.makedirs('output', exist_ok=True)
    converter.document.save(output_file)
    
    print(f'✅ 초기 버전 문서 생성 완료: {output_file}')
    print('📋 이 파일을 열어서 들여쓰기 상태를 확인해주세요.')
    return output_file

if __name__ == "__main__":
    result = generate_initial_version_check()
    print(f"🎉 들여쓰기 확인용 파일 생성: {result}")