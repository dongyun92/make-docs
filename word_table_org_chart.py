#!/usr/bin/env python3
"""
워드 표 기능을 활용한 조직도 생성기
HTML 조직도 대신 네이티브 워드 표로 조직도 생성
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

class WordTableOrgChart:
    def __init__(self):
        self.doc = Document()
        
    def set_cell_border(self, cell, **kwargs):
        """셀 테두리 설정"""
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        
        # 테두리 요소 생성
        tcBorders = tcPr.first_child_found_in(qn('w:tcBorders'))
        if tcBorders is None:
            tcBorders = OxmlElement('w:tcBorders')
            tcPr.append(tcBorders)
        
        # 각 방향별 테두리 설정
        for edge in ('left', 'top', 'right', 'bottom'):
            edge_data = kwargs.get(edge)
            if edge_data:
                tag = f'w:{edge}'
                element = tcBorders.find(qn(tag))
                if element is None:
                    element = OxmlElement(tag)
                    tcBorders.append(element)
                
                element.set(qn('w:val'), edge_data.get('val', 'single'))
                element.set(qn('w:sz'), str(edge_data.get('sz', 4)))
                element.set(qn('w:space'), str(edge_data.get('space', 0)))
                element.set(qn('w:color'), edge_data.get('color', '000000'))

    def create_organization_chart(self):
        """첨단 민군 혁신 지원 시스템 조직도 생성"""
        
        # 제목 추가
        title = self.doc.add_heading('첨단 민군 혁신 지원 시스템 추진 조직도', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 조직도 표 생성 (7행 x 4열)
        table = self.doc.add_table(rows=7, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # 표 너비 설정
        table.width = Inches(6.5)
        
        # 1단계: 사업추진위원회 (최상위, 4개 셀 병합)
        top_cell = table.cell(0, 0)
        for col in range(1, 4):
            top_cell.merge(table.cell(0, col))
        
        top_cell.text = "사업추진위원회\n(총괄책임기관)"
        top_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        top_paragraph = top_cell.paragraphs[0]
        top_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 최상위 셀 스타일링
        self._style_cell(top_cell, bg_color='2C3E50', text_color='FFFFFF', bold=True, font_size=14)
        
        # 연결선 행 (1행)
        for col in range(4):
            connector_cell = table.cell(1, col)
            connector_cell.text = "│" if col == 1 or col == 2 else ""
            connector_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = connector_cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._style_cell(connector_cell, font_size=16, text_color='34495E')
        
        # 2단계: 3개 총괄기관 (2행)
        level2_orgs = [
            ("기술개발\n총괄기관", '3498DB'),
            ("시험평가\n전담기관", '3498DB'), 
            ("사업화지원\n전담기관", '3498DB')
        ]
        
        # 2단계는 1,2,3 컬럼에 배치 (0컬럼은 빈칸)
        table.cell(2, 0).text = ""
        for i, (org_name, color) in enumerate(level2_orgs):
            cell = table.cell(2, i+1)
            cell.text = org_name
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._style_cell(cell, bg_color=color, text_color='FFFFFF', bold=True, font_size=12)
        
        # 연결선 행 (3행)
        for col in range(4):
            connector_cell = table.cell(3, col)
            if col == 1 or col == 2:
                connector_cell.text = "│"
            else:
                connector_cell.text = ""
            connector_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = connector_cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._style_cell(connector_cell, font_size=16, text_color='34495E')
        
        # 3단계: 민간기업 컨소시엄, 정부출연 연구기관 (4행)
        level3_orgs = [
            ("민간기업\n컨소시엄", 'E74C3C'),
            ("정부출연\n연구기관", 'E74C3C')
        ]
        
        table.cell(4, 0).text = ""
        for i, (org_name, color) in enumerate(level3_orgs):
            cell = table.cell(4, i+1)
            cell.text = org_name
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._style_cell(cell, bg_color=color, text_color='FFFFFF', bold=True, font_size=12)
        table.cell(4, 3).text = ""
        
        # 연결선 행 (5행)
        for col in range(4):
            connector_cell = table.cell(5, col)
            connector_cell.text = "│" if col in [1, 2] else ""
            connector_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = connector_cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._style_cell(connector_cell, font_size=16, text_color='34495E')
        
        # 4단계: 4개 개발팀 (6행)
        level4_teams = [
            ("무인체계\n개발팀", '27AE60'),
            ("AI/빅데이터\n개발팀", '27AE60'),
            ("통신기술\n개발팀", '27AE60'), 
            ("센서/반도체\n개발팀", '27AE60')
        ]
        
        for i, (team_name, color) in enumerate(level4_teams):
            cell = table.cell(6, i)
            cell.text = team_name
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._style_cell(cell, bg_color=color, text_color='FFFFFF', bold=True, font_size=11)
        
        # 행 높이 설정
        for i, row in enumerate(table.rows):
            if i in [1, 3, 5]:  # 연결선 행들
                row.height = Inches(0.3)
            else:
                row.height = Inches(0.8)
        
        # 참여기관 정보 추가
        self.doc.add_paragraph()
        participants = self.doc.add_paragraph("참여기관: 대기업 5개사, 중소기업 15개사, 정부출연연 8개 기관, 대학 12개교")
        participants.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 참여기관 텍스트 스타일링
        run = participants.runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(127, 140, 141)  # 회색
        run.italic = True
        
        return table
    
    def _style_cell(self, cell, bg_color=None, text_color='000000', bold=False, font_size=12):
        """셀 스타일링"""
        # 배경색 설정
        if bg_color:
            # 16진수 색상을 RGB로 변환
            rgb_color = tuple(int(bg_color[i:i+2], 16) for i in (0, 2, 4))
            
            # 셀 배경색 설정 (XML 조작)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), bg_color)
            tcPr.append(shd)
        
        # 텍스트 스타일 설정
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = bold
                run.font.size = Pt(font_size)
                
                # 텍스트 색상 설정
                if text_color != '000000':
                    if text_color == 'FFFFFF':
                        run.font.color.rgb = RGBColor(255, 255, 255)
                    elif text_color == '34495E':
                        run.font.color.rgb = RGBColor(52, 73, 94)
                    else:
                        # 16진수 색상을 RGB로 변환
                        rgb_color = tuple(int(text_color[i:i+2], 16) for i in (0, 2, 4))
                        run.font.color.rgb = RGBColor(*rgb_color)
        
        # 셀 패딩 설정
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        
        for margin in ['top', 'left', 'bottom', 'right']:
            mar = OxmlElement(f'w:{margin}')
            mar.set(qn('w:w'), '100')  # 100 twips = ~1.4pt
            mar.set(qn('w:type'), 'dxa')
            tcMar.append(mar)
        
        tcPr.append(tcMar)
        
        # 테두리 설정
        border_style = {
            'val': 'single',
            'sz': 6,
            'space': 0,
            'color': '000000'
        }
        
        self.set_cell_border(
            cell,
            top=border_style,
            left=border_style,
            right=border_style,
            bottom=border_style
        )
    
    def save_document(self, filename="조직도_워드표.docx"):
        """문서 저장"""
        self.doc.save(filename)
        print(f"워드 표 조직도가 {filename}에 저장되었습니다.")
        
    def create_simplified_org_chart(self):
        """간소화된 조직도 (5단계 포함)"""
        
        # 제목
        title = self.doc.add_heading('첨단 민군 혁신 지원 시스템 추진 조직도', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 전체를 위한 큰 표 생성 (11행 x 4열)
        table = self.doc.add_table(rows=11, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 행별 내용 설정
        org_structure = [
            # (행, 병합범위, 내용, 배경색, 연결선여부)
            (0, (0, 4), "사업추진위원회\n(총괄책임기관)", '2C3E50', False),
            (1, None, None, None, True),  # 연결선
            (2, [(1, 1), (2, 1), (3, 1)], ["기술개발\n총괄기관", "시험평가\n전담기관", "사업화지원\n전담기관"], '3498DB', False),
            (3, None, None, None, True),  # 연결선
            (4, [(1, 1), (2, 1)], ["민간기업\n컨소시엄", "정부출연\n연구기관"], 'E74C3C', False),
            (5, None, None, None, True),  # 연결선
            (6, [(0, 1), (1, 1), (2, 1), (3, 1)], ["무인체계\n개발팀", "AI/빅데이터\n개발팀", "통신기술\n개발팀", "센서/반도체\n개발팀"], '27AE60', False),
            (7, None, None, None, True),  # 연결선
            (8, [(1, 1), (2, 1), (3, 1)], ["품질관리팀", "기술사업팀", "국제협력팀"], 'F39C12', False),
        ]
        
        # 표 구성
        for row_idx in range(11):
            for col_idx in range(4):
                cell = table.cell(row_idx, col_idx)
                
                if row_idx in [1, 3, 5, 7]:  # 연결선 행들
                    if col_idx in [1, 2]:
                        cell.text = "│"
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        paragraph = cell.paragraphs[0]
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        self._style_cell(cell, font_size=16, text_color='34495E')
                    else:
                        cell.text = ""
                        
        # Level 1: 사업추진위원회
        level1_cell = table.cell(0, 0)
        for col in range(1, 4):
            level1_cell.merge(table.cell(0, col))
        level1_cell.text = "사업추진위원회\n(총괄책임기관)"
        level1_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        level1_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._style_cell(level1_cell, bg_color='2C3E50', text_color='FFFFFF', bold=True, font_size=14)
        
        # Level 2: 3개 총괄기관
        level2_orgs = ["기술개발\n총괄기관", "시험평가\n전담기관", "사업화지원\n전담기관"]
        table.cell(2, 0).text = ""
        for i, org in enumerate(level2_orgs):
            cell = table.cell(2, i+1)
            cell.text = org
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._style_cell(cell, bg_color='3498DB', text_color='FFFFFF', bold=True, font_size=12)
            
        # Level 3: 2개 기관
        level3_orgs = ["민간기업\n컨소시엄", "정부출연\n연구기관"]
        table.cell(4, 0).text = ""
        table.cell(4, 3).text = ""
        for i, org in enumerate(level3_orgs):
            cell = table.cell(4, i+1)
            cell.text = org
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._style_cell(cell, bg_color='E74C3C', text_color='FFFFFF', bold=True, font_size=12)
        
        # Level 4: 4개 개발팀
        level4_teams = ["무인체계\n개발팀", "AI/빅데이터\n개발팀", "통신기술\n개발팀", "센서/반도체\n개발팀"]
        for i, team in enumerate(level4_teams):
            cell = table.cell(6, i)
            cell.text = team
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._style_cell(cell, bg_color='27AE60', text_color='FFFFFF', bold=True, font_size=11)
        
        # Level 5: 3개 관리팀
        level5_teams = ["품질관리팀", "기술사업팀", "국제협력팀"]
        table.cell(8, 0).text = ""
        for i, team in enumerate(level5_teams):
            cell = table.cell(8, i+1)
            cell.text = team
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._style_cell(cell, bg_color='F39C12', text_color='FFFFFF', bold=True, font_size=11)
        
        # 행 높이 조정
        for i, row in enumerate(table.rows):
            if i in [1, 3, 5, 7]:  # 연결선 행
                row.height = Inches(0.25)
            else:
                row.height = Inches(0.7)
        
        # 참여기관 정보
        self.doc.add_paragraph()
        participants = self.doc.add_paragraph("참여기관: 대기업 5개사, 중소기업 15개사, 정부출연연 8개 기관, 대학 12개교")
        participants.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = participants.runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(127, 140, 141)
        run.italic = True
        
        return table

if __name__ == "__main__":
    chart_creator = WordTableOrgChart()
    
    # 조직도 생성
    print("워드 표 기반 조직도 생성 중...")
    table = chart_creator.create_simplified_org_chart()
    
    # 문서 저장
    chart_creator.save_document("조직도_워드표_완전판.docx")
    
    print("완료! 조직도_워드표_완전판.docx 파일을 확인해보세요.")