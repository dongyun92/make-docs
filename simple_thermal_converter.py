#!/usr/bin/env python3
"""
Thermal Drone Detection MD to DOCX Converter
사용자 요구사항에 맞게 완전히 새로 작성한 변환기
"""

import os
import re
from typing import List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

class SimpleThermalConverter:
    def __init__(self):
        self.document = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """Word 스타일 설정"""
        styles = self.document.styles
        
        # 제목 스타일
        if 'CustomHeading1' not in styles:
            h1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            h1_font = h1_style.font
            h1_font.name = 'Arial'
            h1_font.size = Pt(16)
            h1_font.bold = True
            h1_style.paragraph_format.space_before = Pt(18)
            h1_style.paragraph_format.space_after = Pt(12)
            
        if 'CustomHeading2' not in styles:
            h2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            h2_font = h2_style.font  
            h2_font.name = 'Arial'
            h2_font.size = Pt(14)
            h2_font.bold = True
            h2_style.paragraph_format.space_before = Pt(12)
            h2_style.paragraph_format.space_after = Pt(8)
            
    def convert(self, md_file: str) -> str:
        """메인 변환 함수"""
        print(f"🔄 변환 시작: {md_file}")
        
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
            
        lines = content.split('\n')
        
        # 제목 먼저 추가
        title_para = self.document.add_paragraph("실화상 카메라 기반 장거리 드론탐지시스템 구축 사업계획서")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        
        self.document.add_page_break()
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:  # 빈 줄
                i += 1
                continue
                
            if line.startswith('# '):  # H1 제목 (문서 제목이므로 스킵)
                i += 1
                continue
                
            elif line.startswith('## '):  # H2 제목  
                title = line[3:].strip()
                if title == '주석':  # 주석 섹션은 따로 처리
                    i = self.process_footnote_section(lines, i)
                    continue
                else:
                    para = self.document.add_paragraph(title)
                    para.style = 'CustomHeading1'
                    
            elif line.startswith('### '):  # H3 제목
                title = line[4:].strip()
                para = self.document.add_paragraph(title)
                para.style = 'CustomHeading2'
                
            elif line.startswith('!['):  # 이미지
                i = self.process_image(lines, i)
                continue
                
            elif line.startswith('<그림') or line.startswith('<표'):  # 캡션
                para = self.document.add_paragraph(line)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
                
            elif line.startswith('|') and '|' in line:  # 테이블
                i = self.process_table(lines, i)
                continue
                
            elif line.startswith('□') or line.startswith('○'):  # 불릿 포인트
                self.add_bullet_paragraph(line)
                
            elif line.strip() != '---':  # 일반 텍스트 (구분선 제외)
                if line.strip():  # 빈 줄이 아닌 경우만
                    para = self.document.add_paragraph(line)
                    run = para.runs[0]
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    
            i += 1
            
        # DOCX 저장
        output_file = f"output/{os.path.basename(md_file).replace('.md', '_FIXED.docx')}"
        os.makedirs('output', exist_ok=True)
        self.document.save(output_file)
        
        print(f"✅ 변환 완료: {output_file}")
        return output_file
        
    def process_footnote_section(self, lines: List[str], start_idx: int) -> int:
        """주석 섹션 처리 - 중복 방지"""
        print("📝 주석 섹션 처리 중...")
        
        # 주석 제목 추가 (한 번만!)
        self.document.add_page_break()
        title_para = self.document.add_paragraph("주석")
        title_para.style = 'CustomHeading1'
        
        i = start_idx + 1  # "## 주석" 다음 라인부터
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
                
            if line.startswith('#'):  # 다음 섹션 시작
                break
                
            # 주석 내용 추가
            para = self.document.add_paragraph(line)
            run = para.runs[0]
            run.font.name = 'Arial'  
            run.font.size = Pt(11)  # 동일한 크기로 통일
            
            i += 1
            
        return i
        
    def process_image(self, lines: List[str], start_idx: int) -> int:
        """이미지 처리 - 캡션까지 함께 처리"""
        line = lines[start_idx].strip()
        
        # ![alt](path) 형식 파싱
        match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if match:
            alt_text = match.group(1)
            image_path = match.group(2)
            
            # 절대 경로로 변환
            if not os.path.isabs(image_path):
                current_dir = os.path.dirname(os.path.abspath(__file__))
                full_path = os.path.join(current_dir, image_path)
            else:
                full_path = image_path
                
            print(f"🖼️  이미지 처리: {image_path} -> {full_path}")
            
            if os.path.exists(full_path):
                # 이미지 추가
                para = self.document.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                try:
                    run.add_picture(full_path, width=Inches(5))
                    print(f"✅ 이미지 추가 성공: {image_path}")
                except Exception as e:
                    print(f"❌ 이미지 추가 실패: {e}")
                    # 실패시 텍스트로 표시
                    para = self.document.add_paragraph(f"[이미지: {alt_text}]")
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                print(f"⚠️  이미지 파일 없음: {full_path}")
                # 파일이 없으면 텍스트로 표시
                para = self.document.add_paragraph(f"[이미지 없음: {alt_text} - {image_path}]")
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 다음 줄이 캡션인지 확인하고 함께 처리
            next_idx = start_idx + 1
            if next_idx < len(lines):
                next_line = lines[next_idx].strip()
                if next_line.startswith('<그림') or next_line.startswith('<표'):
                    # 캡션 추가
                    caption_para = self.document.add_paragraph(next_line)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.runs[0]
                    caption_run.font.name = 'Arial'
                    caption_run.font.size = Pt(10)
                    caption_run.font.bold = True
                    print(f"📝 캡션 추가: {next_line}")
                    return next_idx + 1  # 캡션까지 처리했으므로 +2
                
        return start_idx + 1
        
    def process_table(self, lines: List[str], start_idx: int) -> int:
        """테이블 처리"""
        table_lines = []
        i = start_idx
        
        # 테이블 라인들 수집
        while i < len(lines) and lines[i].strip().startswith('|'):
            line = lines[i].strip()
            if not line.startswith('|---'):  # 구분선 제외
                table_lines.append(line)
            i += 1
            
        if len(table_lines) < 2:  # 최소 헤더 + 1행
            return i
            
        # 첫 번째 행에서 열 수 계산
        header_cells = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        col_count = len(header_cells)
        
        if col_count == 0:
            return i
            
        # 테이블 생성
        table = self.document.add_table(rows=1, cols=col_count)
        table.style = 'Table Grid'
        
        # 헤더 추가
        header_row = table.rows[0]
        for j, cell_text in enumerate(header_cells):
            if j < len(header_row.cells):
                header_row.cells[j].text = cell_text
                # 헤더 셀 볼드 처리
                for run in header_row.cells[j].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                    
        # 데이터 행들 추가  
        for line in table_lines[1:]:
            data_cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(data_cells) >= col_count:
                row = table.add_row()
                for j, cell_text in enumerate(data_cells[:col_count]):
                    row.cells[j].text = cell_text
                    # **텍스트** 볼드 처리
                    if '**' in cell_text:
                        cell_text_processed = cell_text.replace('**', '')
                        row.cells[j].text = cell_text_processed
                        for run in row.cells[j].paragraphs[0].runs:
                            run.font.bold = True
                    
                    # 폰트 설정
                    for run in row.cells[j].paragraphs[0].runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                        
        return i
        
    def add_bullet_paragraph(self, text: str):
        """불릿 포인트 단락 추가"""
        para = self.document.add_paragraph(text)
        run = para.runs[0]
        run.font.name = 'Arial'
        run.font.size = Pt(11)


if __name__ == "__main__":
    converter = SimpleThermalConverter()
    result = converter.convert("thermal_drone_detection.md")
    print(f"\n🎉 변환 완료!\n파일: {result}")