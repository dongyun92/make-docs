#!/usr/bin/env python3
"""
MD to DOCX Converter
마크다운 파일을 깔끔한 양식의 워드 문서로 변환하는 도구
"""

import re
import os
import sys
from typing import Dict, List, Tuple, Optional
from pathlib import Path

# 마크다운 처리를 위한 모듈들
# import markdown
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn


class FootnoteManager:
    """워드 주석(footnote) 관리 클래스"""
    
    def __init__(self, document: Document):
        self.document = document
        self.footnote_counter = 0
        self.footnotes = {}  # 주석 번호 -> 주석 내용 매핑
    
    def add_footnote(self, paragraph, text_before_footnote: str, footnote_content: str):
        """실제 워드 주석 추가"""
        self.footnote_counter += 1
        
        # 주석 전 텍스트 추가
        run1 = paragraph.add_run(text_before_footnote)
        run1.font.name = 'Malgun Gothic'
        run1.font.size = Pt(11)
        
        # 주석 참조 추가 (상첨자로)
        footnote_ref = paragraph.add_run(str(self.footnote_counter))
        footnote_ref.font.superscript = True
        footnote_ref.font.size = Pt(8)
        footnote_ref.font.name = 'Malgun Gothic'
        
        # 주석 내용 저장 (나중에 문서 끝에 추가)
        self.footnotes[self.footnote_counter] = footnote_content
        
        return self.footnote_counter
    
    def add_footnotes_section(self):
        """문서 끝에 주석 섹션 추가"""
        if not self.footnotes:
            return
        
        # 주석 섹션 제목
        self.document.add_page_break()
        title_para = self.document.add_paragraph("주석")
        title_para.style = 'CustomHeading1'
        
        # 각 주석 추가
        for num, content in sorted(self.footnotes.items()):
            footnote_para = self.document.add_paragraph()
            
            # 주석 번호 (상첨자)
            num_run = footnote_para.add_run(str(num))
            num_run.font.superscript = True
            num_run.font.size = Pt(8)
            num_run.font.name = 'Malgun Gothic'
            
            # 주석 내용
            content_run = footnote_para.add_run(f" {content}")
            content_run.font.name = 'Malgun Gothic'
            content_run.font.size = Pt(9)
            
            footnote_para.style = 'CustomReference'


class DocumentStyler:
    """DOCX 문서 스타일링을 담당하는 클래스"""
    
    def __init__(self, document: Document):
        self.document = document
        self.setup_styles()
    
    def setup_styles(self):
        """문서 스타일 설정 - 참고 PDF의 전문적인 한국 정부 문서 양식 적용"""
        styles = self.document.styles
        
        # 표지 제목 스타일 (대제목)
        if 'CoverTitle' not in styles:
            cover_title_style = styles.add_style('CoverTitle', WD_STYLE_TYPE.PARAGRAPH)
            cover_title_font = cover_title_style.font
            cover_title_font.name = 'Malgun Gothic'
            cover_title_font.size = Pt(24)
            cover_title_font.bold = True
            cover_title_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            cover_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_title_style.paragraph_format.space_after = Pt(24)
            cover_title_style.paragraph_format.space_before = Pt(60)
        
        # 표지 부제목 스타일
        if 'CoverSubtitle' not in styles:
            cover_subtitle_style = styles.add_style('CoverSubtitle', WD_STYLE_TYPE.PARAGRAPH)
            cover_subtitle_font = cover_subtitle_style.font
            cover_subtitle_font.name = 'Malgun Gothic'
            cover_subtitle_font.size = Pt(18)
            cover_subtitle_font.bold = False
            cover_subtitle_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            cover_subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_subtitle_style.paragraph_format.space_after = Pt(48)
        
        # 표지 정보 스타일 (날짜, 조직 등)
        if 'CoverInfo' not in styles:
            cover_info_style = styles.add_style('CoverInfo', WD_STYLE_TYPE.PARAGRAPH)
            cover_info_font = cover_info_style.font
            cover_info_font.name = 'Malgun Gothic'
            cover_info_font.size = Pt(14)
            cover_info_font.bold = False
            cover_info_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            cover_info_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_info_style.paragraph_format.space_after = Pt(12)
        
        # 문서 제목 스타일 (내용 페이지)
        if 'CustomTitle' not in styles:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Malgun Gothic'
            title_font.size = Pt(20)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 검은색
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(18)
            title_style.paragraph_format.space_before = Pt(12)
        
        # 헤딩 1 스타일 (1. 2. 3... 최상위 번호)
        if 'CustomHeading1' not in styles:
            h1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            h1_font = h1_style.font
            h1_font.name = 'Malgun Gothic'
            h1_font.size = Pt(16)
            h1_font.bold = False
            h1_font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 검은색
            h1_style.paragraph_format.space_before = Pt(18)
            h1_style.paragraph_format.space_after = Pt(12)
            h1_style.paragraph_format.left_indent = Inches(0)  # 들여쓰기 없음
            h1_style.paragraph_format.keep_with_next = True
        
        # 헤딩 2 스타일 (1.1 1.2... 2단계 번호)
        if 'CustomHeading2' not in styles:
            h2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            h2_font = h2_style.font
            h2_font.name = 'Malgun Gothic'
            h2_font.size = Pt(14)
            h2_font.bold = False
            h2_font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 검은색
            h2_style.paragraph_format.space_before = Pt(12)
            h2_style.paragraph_format.space_after = Pt(8)
            h2_style.paragraph_format.left_indent = Inches(0.2)
        
        # 헤딩 3 스타일 (1.1.1 1.1.2... 3단계 번호)
        if 'CustomHeading3' not in styles:
            h3_style = styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
            h3_font = h3_style.font
            h3_font.name = 'Malgun Gothic'
            h3_font.size = Pt(12)
            h3_font.bold = False
            h3_font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 검은색
            h3_style.paragraph_format.space_before = Pt(8)
            h3_style.paragraph_format.space_after = Pt(6)
            h3_style.paragraph_format.left_indent = Inches(0.4)
        
        # 본문 스타일 (표준 단락)
        if 'CustomBody' not in styles:
            body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Malgun Gothic'
            body_font.size = Pt(11)
            body_style.paragraph_format.line_spacing = 1.3
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.first_line_indent = Inches(0.2)
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # 리스트 스타일
        if 'CustomList' not in styles:
            list_style = styles.add_style('CustomList', WD_STYLE_TYPE.PARAGRAPH)
            list_font = list_style.font
            list_font.name = 'Malgun Gothic'
            list_font.size = Pt(11)
            list_style.paragraph_format.left_indent = Inches(0.4)
            try:
                list_style.paragraph_format.hanging_indent = Inches(0.2)
            except AttributeError:
                # hanging_indent 속성이 없는 경우 left_indent만 설정
                list_style.paragraph_format.left_indent = Inches(0.2)
            list_style.paragraph_format.space_after = Pt(3)
            list_style.paragraph_format.line_spacing = 1.2
        
        # 목차 제목 스타일
        if 'TOCHeading' not in styles:
            toc_heading_style = styles.add_style('TOCHeading', WD_STYLE_TYPE.PARAGRAPH)
            toc_heading_font = toc_heading_style.font
            toc_heading_font.name = 'Malgun Gothic'
            toc_heading_font.size = Pt(16)  # 크기 줄임
            toc_heading_font.bold = True
            toc_heading_font.color.rgb = RGBColor(0x00, 0x00, 0x00)  # 검은색
            toc_heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            toc_heading_style.paragraph_format.space_after = Pt(18)
        
        # 목차 항목 스타일 (대제목용)
        if 'TOCEntry1' not in styles:
            toc_entry1_style = styles.add_style('TOCEntry1', WD_STYLE_TYPE.PARAGRAPH)
            toc_entry1_font = toc_entry1_style.font
            toc_entry1_font.name = 'Malgun Gothic'
            toc_entry1_font.size = Pt(12)
            toc_entry1_font.bold = True
            toc_entry1_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            toc_entry1_style.paragraph_format.space_after = Pt(6)
            # 탭 스톱으로 점선 리더 설정
            from docx.enum.text import WD_TAB_ALIGNMENT, WD_TAB_LEADER
            toc_entry1_style.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        
        # 목차 항목 스타일 (중제목용)
        if 'TOCEntry2' not in styles:
            toc_entry2_style = styles.add_style('TOCEntry2', WD_STYLE_TYPE.PARAGRAPH)
            toc_entry2_font = toc_entry2_style.font
            toc_entry2_font.name = 'Malgun Gothic'
            toc_entry2_font.size = Pt(11)
            toc_entry2_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            toc_entry2_style.paragraph_format.space_after = Pt(3)
            toc_entry2_style.paragraph_format.left_indent = Inches(0.3)
            toc_entry2_style.paragraph_format.tab_stops.add_tab_stop(Inches(4.8), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        
        # □ 스타일 1단계 불릿 리스트 (최상위)
        if 'CustomListLevel1' not in styles:
            level1_style = styles.add_style('CustomListLevel1', WD_STYLE_TYPE.PARAGRAPH)
            level1_font = level1_style.font
            level1_font.name = 'Malgun Gothic'
            level1_font.size = Pt(11)
            level1_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            level1_style.paragraph_format.left_indent = Inches(0.25)  # □ 들여쓰기 - 1.1.1(0.2)보다 약간 더
            level1_style.paragraph_format.space_after = Pt(3)
            level1_style.paragraph_format.line_spacing = 1.2
        
        # ○ 스타일 2단계 불릿 리스트 (□보다 더 들여쓰기)
        if 'CustomListLevel2' not in styles:
            level2_style = styles.add_style('CustomListLevel2', WD_STYLE_TYPE.PARAGRAPH)
            level2_font = level2_style.font
            level2_font.name = 'Malgun Gothic'
            level2_font.size = Pt(11)
            level2_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            level2_style.paragraph_format.left_indent = Inches(0.5)  # ○ 들여쓰기 - □보다 더
            level2_style.paragraph_format.space_after = Pt(3)
            level2_style.paragraph_format.line_spacing = 1.2
        
        # - 스타일 3단계 불릿 리스트 (○보다 더 들여쓰기)
        if 'CustomListLevel3' not in styles:
            level3_style = styles.add_style('CustomListLevel3', WD_STYLE_TYPE.PARAGRAPH)
            level3_font = level3_style.font
            level3_font.name = 'Malgun Gothic'
            level3_font.size = Pt(11)
            level3_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            level3_style.paragraph_format.left_indent = Inches(0.75)  # - 들여쓰기 - ○보다 더
            level3_style.paragraph_format.space_after = Pt(3)
            level3_style.paragraph_format.line_spacing = 1.2
        
        # • 스타일 4단계 불릿 리스트 (-보다 더 들여쓰기)
        if 'CustomListLevel4' not in styles:
            level4_style = styles.add_style('CustomListLevel4', WD_STYLE_TYPE.PARAGRAPH)
            level4_font = level4_style.font
            level4_font.name = 'Malgun Gothic'
            level4_font.size = Pt(11)
            level4_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            level4_style.paragraph_format.left_indent = Inches(1.0)  # • 들여쓰기 - -보다 더
            level4_style.paragraph_format.space_after = Pt(3)
            level4_style.paragraph_format.line_spacing = 1.2
        
        # 그림/표 캡션 스타일
        if 'CustomCaption' not in styles:
            caption_style = styles.add_style('CustomCaption', WD_STYLE_TYPE.PARAGRAPH)
            caption_font = caption_style.font
            caption_font.name = 'Malgun Gothic'
            caption_font.size = Pt(10)
            caption_font.bold = True
            caption_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_style.paragraph_format.space_before = Pt(6)
            caption_style.paragraph_format.space_after = Pt(6)
        
        # 참고문헌 항목 스타일
        if 'CustomReference' not in styles:
            ref_style = styles.add_style('CustomReference', WD_STYLE_TYPE.PARAGRAPH)
            ref_font = ref_style.font
            ref_font.name = 'Malgun Gothic'
            ref_font.size = Pt(10)
            ref_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            ref_style.paragraph_format.left_indent = Inches(0.3)
            ref_style.paragraph_format.space_after = Pt(3)
            ref_style.paragraph_format.line_spacing = 1.2
    
    def create_cover_page(self, title: str, subtitle: str = None, date: str = None, organization: str = None):
        """표지 페이지 생성"""
        # 빈 공간 추가 (상단 여백)
        for _ in range(6):
            self.document.add_paragraph()
        
        # 표지 제목
        if title:
            title_para = self.document.add_paragraph(title)
            title_para.style = 'CoverTitle'
        
        # 부제목 (있는 경우)
        if subtitle:
            subtitle_para = self.document.add_paragraph(subtitle)
            subtitle_para.style = 'CoverSubtitle'
        
        # 빈 공간 추가 (중간)
        for _ in range(10):
            self.document.add_paragraph()
        
        # 조직 정보 (여러 줄로 표시)
        if organization:
            org_lines = organization.split('\n')
            for line in org_lines:
                if line.strip():
                    org_para = self.document.add_paragraph(line.strip())
                    org_para.style = 'CoverInfo'
        
        # 빈 공간
        self.document.add_paragraph()
        
        # 날짜 정보
        if date:
            date_para = self.document.add_paragraph(date)
            date_para.style = 'CoverInfo'
    
    def create_table_of_contents(self, headings: List[Tuple[str, int]]):
        """목차 생성"""
        # 목차 제목
        toc_title = self.document.add_paragraph("목  차")
        toc_title.style = 'TOCHeading'
        
        # 목차 항목들
        for heading_text, level in headings:
            toc_entry = self.document.add_paragraph()
            toc_entry.style = 'TOCEntry'
            
            # 들여쓰기 적용
            toc_entry.paragraph_format.left_indent = Inches(0.2 * level)
            
            # 번호와 제목 추가
            toc_entry.add_run(heading_text)
        
        # 페이지 구분
        self.document.add_page_break()


class MarkdownParser:
    """간단한 Markdown 파싱 클래스"""
    
    def __init__(self):
        pass
    
    def parse(self, md_content: str) -> List[str]:
        """Markdown을 라인별로 파싱하여 리스트로 반환"""
        return md_content.strip().split('\n')


class DocxConverter:
    """HTML을 DOCX로 변환하는 메인 클래스"""
    
    def __init__(self):
        self.document = Document()
        self.styler = DocumentStyler(self.document)
        self.parser = MarkdownParser()
        self.footnote_manager = FootnoteManager(self.document)
        self.headings = []  # 목차 생성용 헤딩 수집
        self.section_numbers = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0, 6: 0}  # 섹션 번호 관리
        self.figure_counter = 0  # 그림 번호
        self.table_counter = 0   # 표 번호
        self.processed_captions = set()  # 이미 처리된 캡션 추적
        
        # 페이지 여백 설정 - 참고 문서에 맞춰 조정
        sections = self.document.sections
        for section in sections:
            section.top_margin = Inches(1.2)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1.2)
            section.right_margin = Inches(1)

    def add_caption(self, caption_text: str, caption_type: str = "Figure"):
        """Word의 실제 캡션 기능을 사용하여 캡션 추가"""
        # 캡션 번호 증가
        if caption_type.lower() == "figure" or caption_type.lower() == "그림":
            self.figure_counter += 1
            caption_num = self.figure_counter
            label = "그림"
        else:  # table
            self.table_counter += 1
            caption_num = self.table_counter
            label = "표"
        
        # 캡션 단락 생성
        caption_para = self.document.add_paragraph()
        caption_para.style = 'CustomCaption'
        
        # 캡션 텍스트 구성
        full_caption = f"<{label} {caption_num}> {caption_text}"
        caption_para.add_run(full_caption)
        
        # 캡션을 중앙 정렬
        caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        return caption_para

    def _add_figure_caption(self, alt_text: str):
        """그림 캡션 추가"""
        if alt_text:
            self.add_caption(alt_text, "Figure")

    def _add_table_caption(self, caption_text: str):
        """표 캡션 추가 (표 위에)"""
        if caption_text:
            self.add_caption(caption_text, "Table")
    
    def _add_paragraph_with_formatting(self, text: str):
        """마크다운 포매팅이 포함된 단락 추가"""
        para = self.document.add_paragraph()
        para.style = 'CustomBody'
        
        self._add_formatted_text_to_paragraph(para, text, Pt(11))
    
    def _add_formatted_text_to_paragraph(self, paragraph, text: str, font_size):
        """단락에 마크다운 포매팅이 적용된 텍스트 추가"""
        current_pos = 0
        
        # Bold 패턴 찾기: **text**
        bold_pattern = r'\*\*([^*]+)\*\*'
        # Italic 패턴 찾기: *text* (단, **는 제외)
        italic_pattern = r'(?<!\*)\*([^*]+)\*(?!\*)'
        
        # 모든 포매팅 패턴을 찾고 위치순으로 정렬
        matches = []
        
        # Bold 패턴 찾기
        for match in re.finditer(bold_pattern, text):
            matches.append((match.start(), match.end(), 'bold', match.group(1)))
        
        # Italic 패턴 찾기
        for match in re.finditer(italic_pattern, text):
            matches.append((match.start(), match.end(), 'italic', match.group(1)))
        
        # 위치순으로 정렬
        matches.sort(key=lambda x: x[0])
        
        # 겹치는 패턴 제거 (bold가 우선)
        filtered_matches = []
        for i, (start, end, fmt_type, content) in enumerate(matches):
            overlap = False
            for j, (prev_start, prev_end, prev_type, prev_content) in enumerate(filtered_matches):
                if start < prev_end and end > prev_start:  # 겹침 발생
                    if fmt_type == 'bold':  # bold 우선
                        filtered_matches[j] = (start, end, fmt_type, content)
                    overlap = True
                    break
            if not overlap:
                filtered_matches.append((start, end, fmt_type, content))
        
        # 위치순으로 다시 정렬
        filtered_matches.sort(key=lambda x: x[0])
        
        # 텍스트 처리
        for start, end, fmt_type, content in filtered_matches:
            # 포매팅 이전 텍스트 추가
            before_text = text[current_pos:start]
            if before_text:
                run = paragraph.add_run(before_text)
                run.font.name = 'Malgun Gothic'
                run.font.size = font_size
            
            # 포매팅된 텍스트 추가
            formatted_run = paragraph.add_run(content)
            formatted_run.font.name = 'Malgun Gothic'
            formatted_run.font.size = font_size
            
            if fmt_type == 'bold':
                formatted_run.font.bold = True
            elif fmt_type == 'italic':
                formatted_run.font.italic = True
            
            current_pos = end
        
        # 남은 텍스트 추가
        remaining_text = text[current_pos:]
        if remaining_text:
            run = paragraph.add_run(remaining_text)
            run.font.name = 'Malgun Gothic'
            run.font.size = font_size
    
    def convert_markdown_to_docx(self, md_file_path: str, output_path: str = None) -> str:
        """마크다운 파일을 DOCX로 변환 - 원본 텍스트 직접 처리"""
        
        # 파일 읽기
        with open(md_file_path, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 문서 제목 추출 및 추가
        title = self._extract_title(md_content)
        if title:
            title_para = self.document.add_paragraph(title)
            title_para.style = 'CustomTitle'
            self.document.add_paragraph()  # 빈 줄 추가
        
        # 원본 마크다운 직접 처리
        self._process_markdown_directly(md_content)
        
        # 주석 섹션 추가
        self.footnote_manager.add_footnotes_section()
        
        # 출력 파일 경로 결정
        if output_path is None:
            base_name = Path(md_file_path).stem
            # output 폴더가 없으면 생성
            output_dir = Path("output")
            output_dir.mkdir(exist_ok=True)
            output_path = str(output_dir / f"{base_name}.docx")
        else:
            # 사용자 지정 경로 그대로 사용
            output_dir = Path(output_path).parent
            output_dir.mkdir(parents=True, exist_ok=True)
        
        # 파일 저장
        self.document.save(output_path)
        return output_path
    
    def _extract_title(self, md_content: str) -> Optional[str]:
        """마크다운에서 제목 추출"""
        lines = md_content.split('\n')
        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                return line[2:].strip()
        return None
    
    def _extract_subtitle(self, md_content: str) -> Optional[str]:
        """마크다운에서 부제목 추출 (두 번째 H1 또는 첫 번째 H2)"""
        lines = md_content.split('\n')
        title_found = False
        for line in lines:
            line = line.strip()
            if line.startswith('# '):
                if title_found:  # 두 번째 H1
                    return line[2:].strip()
                title_found = True
            elif line.startswith('## ') and title_found:  # 첫 번째 H2
                return line[3:].strip()
        return None
    
    def _collect_headings(self, lines: List[str]):
        """라인에서 헤딩 정보 수집"""
        headings = []
        
        for line in lines:
            line = line.strip()
            if line.startswith('#'):
                level = 0
                for char in line:
                    if char == '#':
                        level += 1
                    else:
                        break
                title = line[level:].strip()
                if title:  # 빈 제목이 아닌 경우만 추가
                    headings.append((level, title))
        
        return headings
    
    def _update_section_numbers(self, level: int):
        """섹션 번호 업데이트"""
        # 현재 레벨의 번호 증가
        self.section_numbers[level] += 1
        
        # 하위 레벨 번호 초기화
        for l in range(level + 1, 7):
            if l in self.section_numbers:
                self.section_numbers[l] = 0
    
    def _create_numbered_title(self, title: str, level: int) -> str:
        """번호가 포함된 제목 생성"""
        if level == 1:
            return f"{self.section_numbers[1]}. {title}"
        elif level == 2:
            return f"{self.section_numbers[1]}.{self.section_numbers[2]} {title}"
        elif level == 3:
            return f"{self.section_numbers[1]}.{self.section_numbers[2]}.{self.section_numbers[3]} {title}"
        else:
            return title
    
    def _split_into_pages(self, md_content: str) -> List[Dict]:
        """마크다운을 페이지별로 분할"""
        pages = []
        current_page = {"number": None, "title": "", "content": "", "type": "normal"}
        
        lines = md_content.split('\n')
        i = 0
        
        while i < len(lines):
            line = lines[i].strip()
            
            # 페이지 구분자 감지
            if line.startswith('### Page '):
                # 이전 페이지 저장
                if current_page["content"].strip():
                    pages.append(current_page.copy())
                
                # 새 페이지 시작
                page_info = line[9:].strip()  # "### Page " 제거
                
                # 페이지 번호와 제목 분리
                if '.' in page_info:
                    parts = page_info.split('.', 1)
                    page_num = parts[0].strip()
                    page_title = parts[1].strip() if len(parts) > 1 else ""
                else:
                    page_num = page_info
                    page_title = ""
                
                # 페이지 타입 결정
                page_type = self._determine_page_type(page_title)
                
                current_page = {
                    "number": page_num,
                    "title": page_title,
                    "content": "",
                    "type": page_type
                }
            else:
                current_page["content"] += line + '\n'
            
            i += 1
        
        # 마지막 페이지 저장
        if current_page["content"].strip():
            pages.append(current_page)
        
        return pages
    
    def _determine_page_type(self, title: str) -> str:
        """페이지 제목을 기반으로 페이지 타입 결정"""
        title_lower = title.lower()
        
        if '표지' in title:
            return 'cover'
        elif '요약' in title or 'summary' in title_lower:
            return 'summary'
        elif '목차' in title or 'contents' in title_lower:
            return 'toc'
        elif '표·그림' in title or '목록' in title:
            return 'list'
        elif '부록' in title or 'appendix' in title_lower:
            return 'appendix'
        else:
            return 'normal'
    
    def _process_page(self, page_info: Dict):
        """페이지 정보에 따라 적절히 처리"""
        page_type = page_info["type"]
        content = page_info["content"]
        
        if page_type == 'cover':
            self._process_cover_page(content)
        elif page_type == 'summary':
            self._process_summary_page(content, page_info["title"])
        elif page_type == 'toc':
            self._process_toc_page(content)
        elif page_type == 'list':
            self._process_list_page(content, page_info["title"])
        else:
            self._process_normal_page(content, page_info["title"])
        
        # 페이지 구분 (마지막 페이지가 아닌 경우)
        self.document.add_page_break()
    
    def _process_cover_page(self, content: str):
        """표지 페이지 처리"""
        lines = content.strip().split('\n')
        title = ""
        subtitle = ""
        org_info = []
        date_info = ""
        
        for line in lines:
            line = line.strip()
            if not line or line.startswith('**※') or line.startswith('(※'):
                continue
            
            if line.startswith('**') and line.endswith('**'):
                title = line[2:-2]
            elif line.startswith('발주기관'):
                org_info.append(line)
            elif line.startswith('수행기관'):
                org_info.append(line)
            elif line.startswith('작성일자'):
                date_info = line.replace('작성일자 :', '').strip()
        
        # 표지 페이지 생성
        self.styler.create_cover_page(
            title=title,
            subtitle=subtitle,
            date=date_info,
            organization='\n'.join(org_info)
        )
    
    def _process_summary_page(self, content: str, page_title: str):
        """요약문 페이지 처리"""
        # 페이지 제목
        title_para = self.document.add_paragraph(page_title)
        title_para.style = 'CustomTitle'
        
        # HTML로 파싱하여 처리
        soup = self.parser.parse(content)
        self._convert_elements(soup)
    
    def _process_toc_page(self, content: str):
        """목차 페이지 처리"""
        title_para = self.document.add_paragraph("목  차")
        title_para.style = 'TOCHeading'
        
        lines = content.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line or line.startswith('---'):
                continue
            
            # 로마숫자나 번호가 있는 목차 항목 처리
            toc_para = self.document.add_paragraph(line)
            toc_para.style = 'TOCEntry'
    
    def _process_list_page(self, content: str, page_title: str):
        """표·그림 목록 페이지 처리"""
        title_para = self.document.add_paragraph(page_title)
        title_para.style = 'TOCHeading'
        
        lines = content.strip().split('\n')
        for line in lines:
            line = line.strip()
            if not line or line.startswith('---') or line.startswith('('):
                continue
            
            list_para = self.document.add_paragraph(line)
            list_para.style = 'TOCEntry'
    
    def _process_normal_page(self, content: str, page_title: str):
        """일반 페이지 처리"""
        # 페이지 제목이 있으면 추가
        if page_title and page_title not in ['개요']:
            title_para = self.document.add_paragraph(page_title)
            title_para.style = 'CustomHeading1'
        
        # 내용을 섹션별로 분리 (--- 구분자 기준)
        sections = content.split('---')
        
        for i, section in enumerate(sections):
            section = section.strip()
            if not section:
                continue
            
            # 계산근거 및 주해 섹션 처리
            if '계산근거 및 주해' in section:
                self._process_footnote_section(section)
            # 참조·링크 섹션 처리
            elif '참조·링크' in section:
                self._process_reference_section(section)
            # 일반 내용 처리
            else:
                soup = self.parser.parse(section)
                self._convert_elements(soup)
    
    def _process_footnote_section(self, content: str):
        """계산근거 및 주해 섹션 처리"""
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line == '계산근거 및 주해':
                footnote_para = self.document.add_paragraph(line)
                footnote_para.style = 'CustomHeading3'
                footnote_para.paragraph_format.space_before = Pt(12)
            else:
                footnote_para = self.document.add_paragraph(line)
                footnote_para.style = 'CustomBody'
                footnote_para.paragraph_format.left_indent = Inches(0.3)
    
    def _process_reference_section(self, content: str):
        """참조·링크 섹션 처리"""
        lines = content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue
            
            if line == '참조·링크':
                ref_para = self.document.add_paragraph(line)
                ref_para.style = 'CustomHeading3'
                ref_para.paragraph_format.space_before = Pt(12)
            else:
                ref_para = self.document.add_paragraph(line)
                ref_para.style = 'CustomBody'
                ref_para.paragraph_format.left_indent = Inches(0.3)
    
    def _convert_elements(self, soup: BeautifulSoup):
        """HTML 요소들을 DOCX 요소로 변환"""
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table', 'hr']):
            if element.name.startswith('h'):
                self._add_heading(element)
            elif element.name == 'p':
                self._add_paragraph(element)
            elif element.name in ['ul', 'ol']:
                self._add_list(element)
            elif element.name == 'table':
                self._add_table(element)
            elif element.name == 'hr':
                self._add_page_break()
    
    def _convert_elements_intelligently(self, soup: BeautifulSoup):
        """HTML 요소들을 DOCX 요소로 지능적으로 변환 - 자연스러운 문서 구조 유지"""
        
        # 모든 요소를 순차적으로 처리하되, Page 마커는 제거
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'table', 'hr']):
            
            # Page 마커 제거 (### Page N 형태)
            if element.name.startswith('h') and element.get_text().strip().startswith('Page '):
                continue
                
            # HR 구분선도 제거 (불필요한 페이지 나누기 방지)
            if element.name == 'hr':
                continue
                
            if element.name.startswith('h'):
                self._add_intelligent_heading(element)
            elif element.name == 'p':
                self._add_intelligent_paragraph(element)
            elif element.name in ['ul', 'ol']:
                self._add_intelligent_list(element)
            elif element.name == 'table':
                self._add_intelligent_table(element)
    
    def _add_heading(self, element):
        """헤딩 요소 추가"""
        text = element.get_text().strip()
        
        # 헤딩 레벨 추출 (h1, h2, h3 등에서 숫자 부분)
        try:
            level = int(element.name[1:])
        except (ValueError, IndexError):
            level = 1
        
        # 페이지 구분자 처리
        if text.startswith('Page '):
            return  # 페이지 구분은 이미 새로운 구조에서 처리됨
        
        # 로마숫자 또는 특수 번호가 이미 있는 경우 그대로 사용
        if self._has_existing_numbering(text):
            numbered_title = text
        else:
            # 섹션 번호 업데이트
            self._update_section_numbers(level)
            # 번호가 포함된 제목 생성
            numbered_title = self._create_numbered_title(text, level)
        
        # 스타일 매핑
        style_map = {
            1: 'CustomHeading1',
            2: 'CustomHeading2',
            3: 'CustomHeading3'
        }
        
        style = style_map.get(level, 'CustomHeading3')
        para = self.document.add_paragraph(numbered_title)
        para.style = style
    
    def _has_existing_numbering(self, text: str) -> bool:
        """텍스트에 이미 번호가 있는지 확인"""
        # 로마숫자 패턴 확인
        roman_pattern = r'^[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫⅩⅢⅩⅣⅩⅤⅩⅥⅩⅦⅩⅧⅩⅨⅩⅩ]+'
        # 일반 번호 패턴 확인
        number_pattern = r'^\d+[\.\)]'
        
        return bool(re.match(roman_pattern, text) or re.match(number_pattern, text))
    
    def _add_paragraph(self, element):
        """단락 요소 추가"""
        text = element.get_text().strip()
        if text:
            para = self.document.add_paragraph(text)
            para.style = 'CustomBody'
    
    def _add_list(self, element):
        """리스트 요소 추가"""
        for li in element.find_all('li'):
            text = li.get_text().strip()
            if text:
                para = self.document.add_paragraph(f"• {text}")
                para.style = 'CustomListLevel4'
    
    def _add_table(self, element):
        """테이블 요소 추가 - 참고 문서 스타일 적용"""
        rows = element.find_all('tr')
        if not rows:
            return
        
        # 열 수 계산
        max_cols = max(len(row.find_all(['th', 'td'])) for row in rows)
        
        # 테이블 생성
        table = self.document.add_table(rows=len(rows), cols=max_cols)
        table.style = 'Table Grid'
        
        # 테이블 전체 스타일 설정
        table.autofit = False
        
        # 데이터 채우기 및 스타일링
        for i, row in enumerate(rows):
            cells = row.find_all(['th', 'td'])
            for j, cell in enumerate(cells):
                if j < max_cols:
                    cell_obj = table.cell(i, j)
                    cell_obj.text = cell.get_text().strip()
                    
                    # 셀 내부 단락 스타일링
                    for paragraph in cell_obj.paragraphs:
                        # 폰트 설정
                        for run in paragraph.runs:
                            run.font.name = 'Malgun Gothic'
                            run.font.size = Pt(10)
                            
                            # 헤더 행 스타일링
                            if i == 0:
                                run.bold = True
                                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                        
                        # 단락 정렬
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                        
                    # 셀 배경색 (헤더) - 간단히 구현
                    if i == 0:
                        self._set_cell_background(cell_obj, "f2f2f2")
        
        # 테이블 간격 추가
        self.document.add_paragraph()
    
    def _set_cell_background(self, cell, color_hex):
        """셀 배경색 설정"""
        from docx.oxml import parse_xml
        shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(
            'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"',
            color_hex
        ))
        cell._tc.get_or_add_tcPr().append(shading_elm)
    
    def _add_page_break(self):
        """페이지 구분 추가"""
        para = self.document.add_paragraph()
        run = para.runs[0] if para.runs else para.add_run()
        run.add_break(WD_BREAK.PAGE)
    
    def _add_intelligent_heading(self, element):
        """지능적 헤딩 추가 - 실제 문서 구조 파악"""
        text = element.get_text().strip()
        
        # 헤딩 레벨 추출
        try:
            level = int(element.name[1:])
        except (ValueError, IndexError):
            level = 1
        
        # 로마숫자나 특수 번호가 이미 있는 경우 그대로 사용
        if self._has_existing_numbering(text):
            final_title = text
        else:
            final_title = text
        
        # 스타일 매핑 - 실제 사용되는 레벨에 맞춰
        if level <= 2:  # H1, H2는 주요 제목
            style = 'CustomHeading1'
        elif level == 3:  # H3는 중제목
            style = 'CustomHeading2'
        else:  # H4 이상은 소제목
            style = 'CustomHeading3'
        
        para = self.document.add_paragraph(final_title)
        para.style = style
    
    def _add_intelligent_paragraph(self, element):
        """지능적 단락 추가 - 불필요한 줄바꿈 방지"""
        text = element.get_text().strip()
        
        # 빈 내용은 건너뛰기
        if not text:
            return
            
        # 편집 안내나 주석은 건너뛰기
        if text.startswith('**※') or text.startswith('(※'):
            return
        
        # LaTeX 수식 처리
        if '\\[' in text and '\\]' in text:
            text = self._convert_latex_formula(text)
        
        # 주석이 포함된 텍스트 처리
        if '^' in text and self._has_footnote_pattern(text):
            self._add_paragraph_with_footnotes(text)
            return
            
        # 로마숫자로 시작하는 대제목 (Ⅰ, Ⅱ, Ⅲ...)
        if re.match(r'^[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫⅩⅢⅩⅣⅩⅤⅩⅥⅩⅦⅩⅧⅩⅨⅩⅩ]+\s+', text):
            para = self.document.add_paragraph(text)
            para.style = 'CustomHeading1'
        # 숫자로 시작하는 중제목 (1. 2. 3.)
        elif re.match(r'^\d+\.\s+', text):
            para = self.document.add_paragraph(text)
            para.style = 'CustomHeading2'
        # 괄호 숫자로 시작하는 소제목 (1) 2) 3))
        elif re.match(r'^\d+\)\s+', text):
            para = self.document.add_paragraph(text)
            para.style = 'CustomHeading3'
        # 불릿 포인트
        elif text.startswith('•'):
            para = self.document.add_paragraph(text)
            para.style = 'CustomList'
        # 특수 번호 (①, ②, ③)
        elif re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', text):
            para = self.document.add_paragraph(text)
            para.style = 'CustomList'
        # 일반 본문
        else:
            para = self.document.add_paragraph(text)
            para.style = 'CustomBody'
    
    def _add_bullet_paragraph_with_footnotes(self, text: str, style_name: str):
        """bullet point에서 주석이 포함된 단락 추가"""
        para = self.document.add_paragraph()
        
        # 주석 패턴 찾기 (^숫자^[내용] 또는 ^숫자^ 형태)
        footnote_pattern = r'\^(\d+)\^(\[([^\]]*)\])?'
        
        current_pos = 0
        for match in re.finditer(footnote_pattern, text):
            # 주석 전 텍스트 추가
            before_text = text[current_pos:match.start()]
            if before_text:
                run = para.add_run(before_text)
                run.font.name = 'Malgun Gothic'
                run.font.size = Pt(11)
            
            # 주석 번호와 내용 추출
            footnote_num = match.group(1)
            footnote_content = match.group(3) if match.group(3) else f"참조 {footnote_num}"
            
            # 주석 참조 추가 (상첨자)
            footnote_ref = para.add_run(footnote_num)
            footnote_ref.font.superscript = True
            footnote_ref.font.size = Pt(8)
            footnote_ref.font.name = 'Malgun Gothic'
            
            # 주석 내용 저장
            self.footnote_manager.footnotes[int(footnote_num)] = footnote_content
            if int(footnote_num) > self.footnote_manager.footnote_counter:
                self.footnote_manager.footnote_counter = int(footnote_num)
            
            current_pos = match.end()
        
        # 남은 텍스트 추가
        remaining_text = text[current_pos:]
        if remaining_text:
            run = para.add_run(remaining_text)
            run.font.name = 'Malgun Gothic'
            run.font.size = Pt(11)
        
        # 스타일 설정
        para.style = style_name
    
    def _has_footnote_pattern(self, text: str) -> bool:
        """주석 패턴이 있는지 확인 (^숫자^[내용] 또는 ^숫자^ 형태)"""
        return bool(re.search(r'\^\d+\^(\[.*?\])?', text))
    
    def _process_cell_footnotes(self, paragraph, text: str):
        """테이블 셀 내 주석 처리"""
        footnote_pattern = r'\^(\d+)\^(\[([^\]]*)\])?'
        
        current_pos = 0
        for match in re.finditer(footnote_pattern, text):
            # 주석 전 텍스트 추가
            before_text = text[current_pos:match.start()]
            if before_text:
                run = paragraph.add_run(before_text)
                run.font.name = 'Malgun Gothic'
                run.font.size = Pt(10)
            
            # 주석 번호와 내용 추출
            footnote_num = match.group(1)
            footnote_content = match.group(3) if match.group(3) else f"참조 {footnote_num}"
            
            # 주석 참조 추가 (상첨자)
            footnote_ref = paragraph.add_run(footnote_num)
            footnote_ref.font.superscript = True
            footnote_ref.font.size = Pt(8)
            footnote_ref.font.name = 'Malgun Gothic'
            
            # 주석 내용 저장
            self.footnote_manager.footnotes[int(footnote_num)] = footnote_content
            if int(footnote_num) > self.footnote_manager.footnote_counter:
                self.footnote_manager.footnote_counter = int(footnote_num)
            
            current_pos = match.end()
        
        # 남은 텍스트 추가
        remaining_text = text[current_pos:]
        if remaining_text:
            run = paragraph.add_run(remaining_text)
            run.font.name = 'Malgun Gothic'
            run.font.size = Pt(10)

    def _add_paragraph_with_footnotes(self, text: str):
        """주석이 포함된 단락 추가"""
        para = self.document.add_paragraph()
        
        # 주석 패턴 찾기 (^숫자^[내용] 또는 ^숫자^ 형태)
        footnote_pattern = r'\^(\d+)\^(\[([^\]]*)\])?'
        
        current_pos = 0
        for match in re.finditer(footnote_pattern, text):
            # 주석 전 텍스트 추가
            before_text = text[current_pos:match.start()]
            if before_text:
                run = para.add_run(before_text)
                run.font.name = 'Malgun Gothic'
                run.font.size = Pt(11)
            
            # 주석 번호와 내용 추출
            footnote_num = match.group(1)
            footnote_content = match.group(3) if match.group(3) else f"참조 {footnote_num}"
            
            # 주석 참조 추가 (상첨자)
            footnote_ref = para.add_run(footnote_num)
            footnote_ref.font.superscript = True
            footnote_ref.font.size = Pt(8)
            footnote_ref.font.name = 'Malgun Gothic'
            
            # 주석 내용 저장
            self.footnote_manager.footnotes[int(footnote_num)] = footnote_content
            if int(footnote_num) > self.footnote_manager.footnote_counter:
                self.footnote_manager.footnote_counter = int(footnote_num)
            
            current_pos = match.end()
        
        # 남은 텍스트 추가
        remaining_text = text[current_pos:]
        if remaining_text:
            run = para.add_run(remaining_text)
            run.font.name = 'Malgun Gothic'
            run.font.size = Pt(11)
        
        # 스타일 설정
        if self._is_hierarchical_content(text):
            para.style = 'CustomHeading2'
        elif text.startswith(('□', '○', '-', '•')):
            if text.startswith('□'):
                para.style = 'CustomListLevel1'
            elif text.startswith('○'):
                para.style = 'CustomListLevel2'
            elif text.startswith('-'):
                para.style = 'CustomListLevel3'
            elif text.startswith('•'):
                para.style = 'CustomListLevel4'
        else:
            para.style = 'CustomBody'
    
    def _add_hierarchical_paragraph_with_footnotes(self, text: str, style_name: str):
        """계층적 콘텐츠에서 주석이 포함된 단락 추가"""
        para = self.document.add_paragraph()
        
        # 주석 패턴 찾기 (^숫자^[내용] 또는 ^숫자^ 형태)
        footnote_pattern = r'\^(\d+)\^(\[([^\]]*)\])?'
        
        current_pos = 0
        for match in re.finditer(footnote_pattern, text):
            # 주석 전 텍스트 추가
            before_text = text[current_pos:match.start()]
            if before_text:
                run = para.add_run(before_text)
                run.font.name = 'Malgun Gothic'
                if style_name == 'CustomHeading1':
                    run.font.size = Pt(14)
                    run.font.bold = True
                elif style_name == 'CustomHeading2':
                    run.font.size = Pt(13)
                    run.font.bold = True
                elif style_name == 'CustomHeading3':
                    run.font.size = Pt(12)
                    run.font.bold = True
                elif style_name == 'CustomReference':
                    run.font.size = Pt(10)
                else:
                    run.font.size = Pt(11)
            
            # 주석 번호와 내용 추출
            footnote_num = match.group(1)
            footnote_content = match.group(3) if match.group(3) else f"참조 {footnote_num}"
            
            # 주석 참조 추가 (상첨자)
            footnote_ref = para.add_run(footnote_num)
            footnote_ref.font.superscript = True
            footnote_ref.font.size = Pt(8)
            footnote_ref.font.name = 'Malgun Gothic'
            
            # 주석 내용 저장
            self.footnote_manager.footnotes[int(footnote_num)] = footnote_content
            if int(footnote_num) > self.footnote_manager.footnote_counter:
                self.footnote_manager.footnote_counter = int(footnote_num)
            
            current_pos = match.end()
        
        # 남은 텍스트 추가
        remaining_text = text[current_pos:]
        if remaining_text:
            run = para.add_run(remaining_text)
            run.font.name = 'Malgun Gothic'
            if style_name == 'CustomHeading1':
                run.font.size = Pt(14)
                run.font.bold = True
            elif style_name == 'CustomHeading2':
                run.font.size = Pt(13)
                run.font.bold = True
            elif style_name == 'CustomHeading3':
                run.font.size = Pt(12)
                run.font.bold = True
            elif style_name == 'CustomReference':
                run.font.size = Pt(10)
            else:
                run.font.size = Pt(11)
        
        # 스타일 설정
        para.style = style_name
    
    def _add_intelligent_list(self, element):
        """지능적 리스트 추가"""
        for li in element.find_all('li'):
            text = li.get_text().strip()
            if text:
                # 이미 불릿이 있는지 확인
                if not text.startswith('•'):
                    text = f"• {text}"
                para = self.document.add_paragraph(text)
                para.style = 'CustomListLevel4'
    
    def _add_intelligent_table(self, element):
        """지능적 테이블 추가 - 기존 테이블 로직 재사용하되 개선"""
        self._add_table(element)
    
    def _process_markdown_directly(self, md_content: str):
        """원본 마크다운을 직접 라인별로 처리"""
        lines = md_content.split('\n')
        i = 0
        current_table = []
        in_table = False
        
        while i < len(lines):
            original_line = lines[i] 
            line = lines[i].strip()
            
            # 빈 줄 건너뛰기
            if not line:
                i += 1
                continue
            
            # 페이지 마커 건너뛰기
            if line.startswith('### Page '):
                i += 1
                continue
                
            # 편집 안내 건너뛰기
            if line.startswith('**※') or line.startswith('(※'):
                i += 1
                continue
                
            # 구분선 건너뛰기
            if line == '---':
                i += 1
                continue
            
            # 이미지 처리
            if self._is_image_line(line):
                self._add_image(line)
                i += 1
                continue
            
            # 테이블 처리
            if '|' in line and not in_table:
                # 테이블 시작 전에 이전 라인이 표 캡션인지 확인
                table_caption_before = None
                if i > 0:
                    prev_line = lines[i - 1].strip()
                    if re.match(r'^<표\s*\d+>\s*.*', prev_line, re.IGNORECASE):
                        table_caption_before = prev_line
                        print(f"DEBUG: 테이블 위 캡션 발견: {table_caption_before}")
                        
                        # 캡션이 아직 처리되지 않았으면 처리
                        if table_caption_before not in self.processed_captions:
                            match = re.match(r'^<표\s*\d+>\s*(.*)', table_caption_before, re.IGNORECASE)
                            if match:
                                caption_text = match.group(1).strip()
                                self.add_caption(caption_text, "Table")
                                self.processed_captions.add(table_caption_before)
                                print(f"DEBUG: 테이블 위 캡션 처리함: {table_caption_before}")
                        else:
                            print(f"DEBUG: 테이블 위 캡션 이미 처리됨: {table_caption_before}")
                
                # 테이블 시작
                in_table = True
                current_table = [line]
                i += 1
                continue
            elif in_table:
                if '|' in line:
                    current_table.append(line)
                    i += 1
                    continue
                else:
                    # 테이블 끝 - 다음 최대 3줄에서 표 캡션 찾기
                    table_caption = None
                    caption_line_index = None
                    
                    for j in range(1, 4):  # 최대 3줄까지 확인
                        if i + j < len(lines):
                            check_line = lines[i + j].strip()
                            if not check_line:  # 빈 라인은 건너뛰기
                                continue
                            if re.match(r'^<표\s*\d+>\s*.*', check_line, re.IGNORECASE):
                                table_caption = check_line
                                caption_line_index = i + j
                                print(f"DEBUG: 테이블 캡션 발견: {table_caption}")
                                break
                            else:
                                # 다른 텍스트가 나오면 캡션 찾기 중단
                                break
                    
                    # 표 캡션이 있으면 먼저 추가
                    if table_caption:
                        match = re.match(r'^<표\s*\d+>\s*(.*)', table_caption, re.IGNORECASE)
                        if match:
                            caption_text = match.group(1).strip()
                            self.add_caption(caption_text, "Table")
                            self.processed_captions.add(table_caption)  # 처리된 캡션 기록
                            print(f"DEBUG: 표 캡션 처리함: {table_caption}")
                            
                            # 캡션 라인을 건너뛰도록 인덱스 조정
                            if caption_line_index:
                                # 캡션 라인까지 건너뛰기 위해 i 설정
                                i = caption_line_index
                    
                    # 테이블 생성
                    self._create_table_from_lines(current_table)
                    current_table = []
                    in_table = False
                    # 현재 라인은 다시 처리
                    
            # 헤딩 처리 (# ## ###)
            if line.startswith('#'):
                self._process_heading_line(line)
            # 목차 항목 처리 (점선과 페이지 번호가 있는 경우) - 우선 처리
            elif self._is_toc_entry(original_line):
                self._add_toc_entry(original_line)
            # 계층적 번호 체계 처리
            elif self._is_hierarchical_content(line):
                self._add_hierarchical_content(line)
            # 로마숫자로 시작하는 대제목 (Ⅰ, Ⅱ, Ⅲ...)
            elif re.match(r'^[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫⅩⅢⅩⅣⅩⅤⅩⅥⅩⅦⅩⅧⅩⅨⅩⅩ]+\s+', line):
                para = self.document.add_paragraph(line)
                para.style = 'CustomHeading1'
            # 숫자로 시작하는 중제목 (1. 2. 3.) - 계층적 번호 체계가 우선 처리되므로 이 부분은 삭제
            # 괄호 숫자로 시작하는 소제목 (1) 2) 3))
            elif re.match(r'^\d+\)\s+', line):
                para = self.document.add_paragraph(line)
                para.style = 'CustomHeading3'
            # 불릿 포인트 처리
            elif line.startswith('□'):
                if '^' in line and self._has_footnote_pattern(line):
                    self._add_bullet_paragraph_with_footnotes(line, 'CustomListLevel1')
                else:
                    para = self.document.add_paragraph(line)
                    para.style = 'CustomListLevel1'
            elif line.startswith('○'):
                if '^' in line and self._has_footnote_pattern(line):
                    self._add_bullet_paragraph_with_footnotes(line, 'CustomListLevel2')
                else:
                    para = self.document.add_paragraph(line)
                    para.style = 'CustomListLevel2'  
            elif line.startswith('-'):
                if '^' in line and self._has_footnote_pattern(line):
                    self._add_bullet_paragraph_with_footnotes(line, 'CustomListLevel3')
                else:
                    para = self.document.add_paragraph(line)
                    para.style = 'CustomListLevel3'
            elif line.startswith('•'):
                if '^' in line and self._has_footnote_pattern(line):
                    self._add_bullet_paragraph_with_footnotes(line, 'CustomListLevel4')
                else:
                    para = self.document.add_paragraph(line)
                    para.style = 'CustomListLevel4'
            # 특수 번호 (①②③)
            elif re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', line):
                para = self.document.add_paragraph(line)
                para.style = 'CustomList'
            # 일반 텍스트
            else:
                # LaTeX 수식 처리
                if '\\[' in line and '\\]' in line:
                    line = self._convert_latex_formula(line)
                
                # 주석이 포함된 텍스트 처리
                if '^' in line and self._has_footnote_pattern(line):
                    self._add_paragraph_with_footnotes(line)
                else:
                    # 마크다운 포매팅이 있는지 확인하고 처리
                    if '**' in line or '*' in line:
                        self._add_paragraph_with_formatting(line)
                    else:
                        para = self.document.add_paragraph(line)
                        para.style = 'CustomBody'
            
            i += 1
        
        # 마지막에 테이블이 있으면 처리
        if current_table:
            self._create_table_from_lines(current_table)
    
    def _process_heading_line(self, line: str):
        """헤딩 라인 처리"""
        if line.startswith('###'):
            text = line[3:].strip()
            para = self.document.add_paragraph(text)
            para.style = 'CustomHeading3'
        elif line.startswith('##'):
            text = line[2:].strip()
            para = self.document.add_paragraph(text)
            para.style = 'CustomHeading2'
        elif line.startswith('#'):
            text = line[1:].strip()
            para = self.document.add_paragraph(text)
            para.style = 'CustomHeading1'
    
    def _is_hierarchical_content(self, line: str) -> bool:
        """계층적 번호 체계 콘텐츠인지 확인 - 숫자 기반 계층만 처리"""
        # 1. 형태의 번호 (1단계 - 가장 상위)
        if re.match(r'^\d+\.\s+', line):
            return True
        # 1.1 형태의 번호 (2단계)
        if re.match(r'^\d+\.\d+\s+', line):
            return True
        # 1.1.1 형태의 번호 (3단계)
        if re.match(r'^\d+\.\d+\.\d+\s+', line):
            return True
        # 그림/표 캡션 패턴
        if re.match(r'^[\<\[]*\s*(그림|표|figure|table)\s*[\>\]]*\s*\d+', line, re.IGNORECASE):
            return True
        return False
    
    def _add_hierarchical_content(self, line: str):
        """계층적 번호 체계 콘텐츠 추가"""
        # 주석이 포함된 경우 특별 처리
        if '^' in line and self._has_footnote_pattern(line):
            # 1.1.1 형태 - 2단계 헤딩과 동일 (가장 구체적인 패턴부터 확인)
            if re.match(r'^\d+\.\d+\.\d+\s+', line):
                self._add_hierarchical_paragraph_with_footnotes(line, 'CustomHeading2')
            # 1.1 형태 - 2단계 헤딩
            elif re.match(r'^\d+\.\d+\s+', line):
                self._add_hierarchical_paragraph_with_footnotes(line, 'CustomHeading2')
            # 1. 형태 - 최상위 헤딩이지만 참고문헌 항목일 수 있음
            elif re.match(r'^\d+\.\s+', line):
                if self._is_reference_item(line):
                    self._add_hierarchical_paragraph_with_footnotes(line, 'CustomReference')
                else:
                    self._add_hierarchical_paragraph_with_footnotes(line, 'CustomHeading1')
            else:
                self._add_paragraph_with_footnotes(line)
            return
        
        # 일반 처리 (주석 없는 경우)
        # 1.1.1 형태 - 2단계 헤딩과 동일 (가장 구체적인 패턴부터 확인)
        if re.match(r'^\d+\.\d+\.\d+\s+', line):
            para = self.document.add_paragraph(line)
            para.style = 'CustomHeading2'
        # 1.1 형태 - 2단계 헤딩
        elif re.match(r'^\d+\.\d+\s+', line):
            para = self.document.add_paragraph(line)
            para.style = 'CustomHeading2'
        # 1. 형태 - 최상위 헤딩이지만 참고문헌 항목일 수 있음
        elif re.match(r'^\d+\.\s+', line):
            # 참고문헌 항목인지 확인 (저자명, 연도, 제목 등이 포함된 경우)
            if self._is_reference_item(line):
                para = self.document.add_paragraph(line)
                para.style = 'CustomReference'
            else:
                para = self.document.add_paragraph(line)
                para.style = 'CustomHeading1'
        # 그림/표 캡션 - Word 실제 캡션 기능 사용
        elif re.match(r'^[\<\[]*\s*(그림|표|figure|table)\s*[\>\]]*\s*\d+', line, re.IGNORECASE):
            # 이미 처리된 캡션인지 확인
            if line in self.processed_captions:
                print(f"DEBUG: 이미 처리된 캡션 건너뛰기: {line}")
                pass  # 이미 처리된 캡션은 건너뛰기
            else:
                # 캡션 텍스트에서 실제 설명 부분 추출
                match = re.match(r'^<(그림|표|figure|table)\s*\d+>\s*(.*)', line, re.IGNORECASE)
                if match:
                    caption_type = match.group(1)
                    caption_text = match.group(2).strip()
                    
                    if caption_type.lower() in ['그림', 'figure']:
                        self.add_caption(caption_text, "Figure")
                        self.processed_captions.add(line)
                        print(f"DEBUG: 그림 캡션 처리함: {line}")
                    elif caption_type.lower() in ['표', 'table']:
                        # 표 캡션은 테이블과 함께 처리되어야 하므로 여기서는 일반적으로 건너뛰기
                        # 단, 테이블 없이 단독으로 나타나는 경우만 처리
                        print(f"DEBUG: 표 캡션 발견했지만 테이블 처리에서 담당: {line}")
                else:
                    # 패턴이 복잡하면 기존 방식 사용
                    para = self.document.add_paragraph(line)
                    para.style = 'CustomCaption'
        else:
            # 기본 처리
            para = self.document.add_paragraph(line)
            para.style = 'CustomBody'
    
    def _is_reference_item(self, line: str) -> bool:
        """참고문헌 항목인지 확인"""
        # 저자명, 연도, 제목 등이 포함된 참고문헌 패턴
        if re.search(r'\(\d{4}\)', line):  # 연도 패턴
            return True
        if '저자명' in line or '논문제목' in line or '저널명' in line or '출판사명' in line or '보고서제목' in line or '기관명' in line:
            return True
        return False

    def _create_table_from_lines(self, table_lines: list):
        """마크다운 테이블 라인들을 DOCX 테이블로 변환"""
        if not table_lines:
            return
            
        # 구분선 제거
        clean_lines = [line for line in table_lines if not re.match(r'^[\|\-\s]*$', line)]
        
        if not clean_lines:
            return
        
        # 첫 번째 라인으로 열 수 계산
        first_row = [cell.strip() for cell in clean_lines[0].split('|') if cell.strip()]
        cols = len(first_row)
        
        if cols == 0:
            return
            
        # 테이블 생성
        table = self.document.add_table(rows=len(clean_lines), cols=cols)
        table.style = 'Table Grid'
        
        # 데이터 채우기
        for i, line in enumerate(clean_lines):
            cells = [cell.strip() for cell in line.split('|') if cell.strip()]
            for j, cell_text in enumerate(cells):
                if j < cols:
                    cell_obj = table.cell(i, j)
                    
                    # 셀 내용 클리어하고 새로 작성
                    cell_obj.text = ""
                    paragraph = cell_obj.paragraphs[0]
                    
                    # 상첨자 처리가 필요한 텍스트인지 확인
                    if self._has_footnote_pattern(cell_text):
                        self._process_cell_footnotes(paragraph, cell_text)
                    else:
                        run = paragraph.add_run(cell_text)
                        run.font.name = 'Malgun Gothic'
                        run.font.size = Pt(10)
                        if i == 0:  # 헤더 행
                            run.bold = True
                    
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
                    
                    # 헤더 배경색
                    if i == 0:
                        self._set_cell_background(cell_obj, "f2f2f2")
        
        # 테이블 후 간격
        self.document.add_paragraph()
    
    def _convert_latex_formula(self, text: str) -> str:
        """LaTeX 수식을 워드에서 읽기 쉬운 형태로 변환"""
        
        # LaTeX 수식 패턴 찾기
        formula_pattern = r'\\?\[(.*?)\\?\]'
        
        def replace_formula(match):
            formula = match.group(1)
            
            # 기본 LaTeX 명령어들을 유니코드로 변환
            replacements = {
                r'\\text\{(.*?)\}': r'\1',  # \text{} 제거
                r'\\frac\{(.*?)\}\{(.*?)\}': r'(\1)/(\2)',  # 분수를 (분자)/(분모) 형태로
                r'\\times': '×',  # 곱하기
                r'\\div': '÷',    # 나누기
                r'\\pm': '±',     # 플러스마이너스
                r'\\approx': '≈', # 근사값
                r'\\leq': '≤',    # 작거나 같음
                r'\\geq': '≥',    # 크거나 같음
                r'\\alpha': 'α',  # 그리스 문자들
                r'\\beta': 'β',
                r'\\gamma': 'γ',
                r'\\delta': 'δ',
                r'\\epsilon': 'ε',
                r'\\lambda': 'λ',
                r'\\mu': 'μ',
                r'\\pi': 'π',
                r'\\sigma': 'σ',
                r'\\sum': 'Σ',
                r'\\\%': '%',     # 퍼센트
            }
            
            # 각 치환 적용
            for pattern, replacement in replacements.items():
                formula = re.sub(pattern, replacement, formula)
            
            return f"[{formula}]"
        
        # 수식 변환 적용
        converted = re.sub(formula_pattern, replace_formula, text)
        return converted
    
    def _is_toc_entry(self, line: str) -> bool:
        """목차 항목인지 확인 (점선과 숫자가 있는 경우)"""
        # 점선과 마지막에 숫자가 있는 패턴
        return bool(re.search(r'[…\.]{3,}.*?\d+\s*$', line))
    
    def _add_toc_entry(self, line: str):
        """목차 항목 추가"""
        # 점선 패턴 찾기
        match = re.match(r'^(.+?)\s+[…\.]{3,}.*?(\d+)\s*$', line)
        if match:
            title = match.group(1).strip()
            page_num = match.group(2).strip()
            
            # 들여쓰기 레벨 결정 (유니코드 공백 포함)
            if line.startswith(' ') or line.startswith('\u2003') or line.startswith('\t'):
                # 들여쓰기가 있으면 중제목
                style = 'TOCEntry2'
            else:
                # 들여쓰기 없으면 대제목
                style = 'TOCEntry1'
            
            # 목차 항목 생성 (탭으로 구분)
            para = self.document.add_paragraph()
            para.style = style
            para.add_run(title)
            para.add_run('\t')  # 탭 문자로 점선 리더 활성화
            para.add_run(page_num)
    
    def _is_image_line(self, line: str) -> bool:
        """이미지 마크다운 라인인지 확인 ![alt](path)"""
        return bool(re.match(r'!\[.*?\]\(.*?\)', line))
    
    def _add_image(self, line: str):
        """이미지 추가 처리"""
        # 마크다운 이미지 패턴 매칭 ![alt](path)
        match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if not match:
            return
            
        alt_text = match.group(1)
        image_path = match.group(2)
        
        # 상대 경로를 절대 경로로 변환
        if not os.path.isabs(image_path):
            current_dir = os.path.dirname(os.path.abspath(__file__))
            image_path = os.path.join(current_dir, image_path)
        
        # 이미지 파일 존재 확인
        if not os.path.exists(image_path):
            print(f"경고: 이미지 파일을 찾을 수 없습니다: {image_path}")
            # 이미지가 없으면 캡션만 추가
            if alt_text:
                self._add_figure_caption(alt_text)
            return
        
        try:
            # 이미지를 문서에 추가
            paragraph = self.document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 이미지 크기 조정 (최대 너비 6인치)
            run = paragraph.add_run()
            run.add_picture(image_path, width=Inches(6))
            
            # 캡션 추가 (있는 경우) - Word 실제 캡션 기능 사용
            if alt_text:
                self._add_figure_caption(alt_text)
            
            # 이미지 후 간격
            self.document.add_paragraph()
            
        except Exception as e:
            print(f"이미지 삽입 중 오류 발생: {e}")
            # 오류 시 캡션만 추가
            para = self.document.add_paragraph(alt_text)
            para.style = 'CustomCaption'


def main():
    """메인 함수"""
    if len(sys.argv) < 2:
        print("사용법: python md_to_docx_converter.py <input_file.md> [output_file.docx]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    
    if not os.path.exists(input_file):
        print(f"입력 파일을 찾을 수 없습니다: {input_file}")
        sys.exit(1)
    
    try:
        converter = DocxConverter()
        output_path = converter.convert_markdown_to_docx(input_file, output_file)
        print(f"변환 완료: {output_path}")
    except Exception as e:
        import traceback
        print(f"변환 중 오류 발생: {e}")
        print("상세 오류 정보:")
        traceback.print_exc()
        sys.exit(1)


if __name__ == "__main__":
    main()