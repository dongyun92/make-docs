#!/usr/bin/env python3
"""
간단한 워드 표 기반 조직도 생성기
복잡한 XML 조작 없이 기본 기능만 사용
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH

class SimpleWordOrgChart:
    def __init__(self):
        self.doc = Document()
        
    def create_organization_chart(self):
        """간단한 조직도 생성"""
        
        # 제목 추가
        title = self.doc.add_heading('첨단 민군 혁신 지원 시스템 추진 조직도', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 조직도를 여러 개의 작은 표로 분할하여 생성
        
        # Level 1: 최상위 (1x1 표)
        self._create_level_table("사업추진위원회\n(총괄책임기관)", 1, 1, '2C3E50')
        
        # 연결 표시
        self._add_connector()
        
        # Level 2: 3개 총괄기관 (1x3 표)
        level2_data = ["기술개발\n총괄기관", "시험평가\n전담기관", "사업화지원\n전담기관"]
        self._create_level_table(level2_data, 1, 3, '3498DB')
        
        # 연결 표시
        self._add_connector()
        
        # Level 3: 2개 기관 (1x2 표)
        level3_data = ["민간기업 컨소시엄", "정부출연 연구기관"]
        self._create_level_table(level3_data, 1, 2, 'E74C3C')
        
        # 연결 표시
        self._add_connector()
        
        # Level 4: 4개 개발팀 (1x4 표)
        level4_data = ["무인체계\n개발팀", "AI/빅데이터\n개발팀", "통신기술\n개발팀", "센서/반도체\n개발팀"]
        self._create_level_table(level4_data, 1, 4, '27AE60')
        
        # 연결 표시
        self._add_connector()
        
        # Level 5: 3개 관리팀 (1x3 표)
        level5_data = ["품질관리팀", "기술사업팀", "국제협력팀"]
        self._create_level_table(level5_data, 1, 3, 'F39C12')
        
        # 참여기관 정보 추가
        self.doc.add_paragraph()
        participants = self.doc.add_paragraph("참여기관: 대기업 5개사, 중소기업 15개사, 정부출연연 8개 기관, 대학 12개교")
        participants.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 참여기관 텍스트 스타일링
        run = participants.runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(127, 140, 141)  # 회색
        run.italic = True
        
        print("간단한 조직도가 생성되었습니다.")
        
    def _create_level_table(self, data, rows, cols, color_hex):
        """각 레벨별 표 생성"""
        table = self.doc.add_table(rows=rows, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 데이터가 문자열이면 단일 셀 처리
        if isinstance(data, str):
            cell = table.cell(0, 0)
            cell.text = data
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            # 텍스트 스타일링
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 폰트 설정
            run = paragraph.runs[0]
            run.font.bold = True
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(255, 255, 255)  # 흰색 텍스트
            
        # 데이터가 리스트면 여러 셀 처리
        elif isinstance(data, list):
            for i, text in enumerate(data):
                if i < cols:
                    cell = table.cell(0, i)
                    cell.text = text
                    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    
                    # 텍스트 스타일링
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    # 폰트 설정
                    run = paragraph.runs[0]
                    run.font.bold = True
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(255, 255, 255)  # 흰색 텍스트
        
        # 표 스타일 설정 (배경색은 Word에서 수동으로 설정 필요)
        table.style = 'Table Grid'
        
        # 행 높이 설정
        for row in table.rows:
            row.height = Inches(0.8)
        
        # 표 너비 설정
        table.width = Inches(6)
        
        return table
    
    def _add_connector(self):
        """레벨 간 연결선 표시"""
        connector_para = self.doc.add_paragraph("│")
        connector_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 연결선 스타일링
        run = connector_para.runs[0]
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(52, 73, 94)  # 진한 회색
    
    def create_unified_org_chart(self):
        """통합된 단일 표 조직도 (간소화 버전)"""
        
        # 제목 추가
        title = self.doc.add_heading('첨단 민군 혁신 지원 시스템 추진 조직도', level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 9행 5열 표 생성 (연결선 포함)
        table = self.doc.add_table(rows=9, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        # Level 1: 사업추진위원회 (4개 셀 병합)
        top_cell = table.cell(0, 0)
        for col in range(1, 4):
            top_cell.merge(table.cell(0, col))
        
        self._set_cell_content(top_cell, "사업추진위원회\n(총괄책임기관)", 14, True)
        
        # 연결선 1
        for col in range(4):
            cell = table.cell(1, col)
            if col in [1, 2]:
                cell.text = "│"
                self._set_cell_content(cell, "│", 16, False, center_only=True)
            else:
                cell.text = ""
        
        # Level 2: 3개 총괄기관
        table.cell(2, 0).text = ""  # 빈 셀
        orgs = ["기술개발\n총괄기관", "시험평가\n전담기관", "사업화지원\n전담기관"]
        for i, org in enumerate(orgs):
            if i < 3:  # 3개만
                cell = table.cell(2, i + 1 if i < 1 else i)
                self._set_cell_content(cell, org, 12, True)
        
        # 연결선 2
        for col in range(4):
            cell = table.cell(3, col)
            if col in [1, 2]:
                cell.text = "│"
                self._set_cell_content(cell, "│", 16, False, center_only=True)
            else:
                cell.text = ""
        
        # Level 3: 2개 기관
        table.cell(4, 0).text = ""  # 빈 셀
        table.cell(4, 3).text = ""  # 빈 셀
        level3_orgs = ["민간기업\n컨소시엄", "정부출연\n연구기관"]
        for i, org in enumerate(level3_orgs):
            cell = table.cell(4, i + 1)
            self._set_cell_content(cell, org, 12, True)
        
        # 연결선 3
        for col in range(4):
            cell = table.cell(5, col)
            if col in [0, 1, 2, 3]:
                cell.text = "│"
                self._set_cell_content(cell, "│", 16, False, center_only=True)
            else:
                cell.text = ""
        
        # Level 4: 4개 개발팀
        level4_teams = ["무인체계\n개발팀", "AI/빅데이터\n개발팀", "통신기술\n개발팀", "센서/반도체\n개발팀"]
        for i, team in enumerate(level4_teams):
            cell = table.cell(6, i)
            self._set_cell_content(cell, team, 11, True)
        
        # 연결선 4
        for col in range(4):
            cell = table.cell(7, col)
            if col in [1, 2]:
                cell.text = "│"
                self._set_cell_content(cell, "│", 16, False, center_only=True)
            else:
                cell.text = ""
        
        # Level 5: 3개 관리팀
        table.cell(8, 0).text = ""  # 빈 셀
        level5_teams = ["품질관리팀", "기술사업팀", "국제협력팀"]
        for i, team in enumerate(level5_teams):
            cell = table.cell(8, i + 1)
            self._set_cell_content(cell, team, 11, True)
        
        # 행 높이 설정
        for i, row in enumerate(table.rows):
            if i in [1, 3, 5, 7]:  # 연결선 행들
                row.height = Inches(0.3)
            else:
                row.height = Inches(0.8)
        
        # 참여기관 정보
        self.doc.add_paragraph()
        participants = self.doc.add_paragraph("참여기관: 대기업 5개사, 중소기업 15개사, 정부출연연 8개 기관, 대학 12개교")
        participants.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        run = participants.runs[0]
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(127, 140, 141)
        run.italic = True
        
        return table
    
    def _set_cell_content(self, cell, text, font_size, bold, center_only=False):
        """셀 내용 및 스타일 설정"""
        cell.text = text
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # 문단 정렬
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 폰트 설정 (기존 텍스트가 있을 때만)
        if paragraph.runs:
            run = paragraph.runs[0]
            run.font.size = Pt(font_size)
            run.font.bold = bold
            
            if not center_only:  # 연결선이 아닌 경우만 흰색 텍스트
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:  # 연결선은 진한 회색
                run.font.color.rgb = RGBColor(52, 73, 94)
    
    def save_document(self, filename="조직도_간단_워드표.docx"):
        """문서 저장"""
        self.doc.save(filename)
        print(f"워드 표 조직도가 {filename}에 저장되었습니다.")

if __name__ == "__main__":
    chart_creator = SimpleWordOrgChart()
    
    print("간단한 워드 표 기반 조직도 생성 중...")
    
    # 통합된 조직도 생성
    chart_creator.create_unified_org_chart()
    
    # 문서 저장
    chart_creator.save_document("조직도_워드표_간단버전.docx")
    
    print("완료! 조직도_워드표_간단버전.docx 파일을 확인해보세요.")
    print("\n참고: Word에서 표 셀의 배경색은 다음과 같이 설정할 수 있습니다:")
    print("1. 표 전체 선택")
    print("2. 표 디자인 탭 → 음영 → 원하는 색상 선택")
    print("3. 또는 각 셀별로 우클릭 → 표 속성 → 테두리 및 음영에서 설정")