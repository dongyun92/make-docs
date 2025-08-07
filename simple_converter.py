#!/usr/bin/env python3
"""
간단한 MD to DOCX 변환기 - 완전히 새로 작성
"""

import re
import os
import sys
from typing import Dict, List, Tuple, Optional
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement, qn


class SimpleDocxConverter:
    """간단한 DOCX 변환기"""
    
    def __init__(self):
        self.document = Document()
        self.setup_styles()
        self.figure_counter = 0
        self.table_counter = 0
        self.processed_captions = set()
    
    def setup_styles(self):
        """문서 스타일 설정 - 전문적인 한국 정부 문서 양식"""
        styles = self.document.styles
        
        # 표지 제목 스타일 (대제목)
        if 'CoverTitle' not in styles:
            cover_title_style = styles.add_style('CoverTitle', WD_STYLE_TYPE.PARAGRAPH)
            cover_title_font = cover_title_style.font
            cover_title_font.name = 'Malgun Gothic'
            cover_title_font.size = Pt(24)
            cover_title_font.bold = True
            cover_title_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            cover_title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cover_title_style.paragraph_format.space_after = Pt(24)
            cover_title_style.paragraph_format.space_before = Pt(60)
        
        # 문서 제목 스타일 (내용 페이지)
        if 'CustomTitle' not in styles:
            title_style = styles.add_style('CustomTitle', WD_STYLE_TYPE.PARAGRAPH)
            title_font = title_style.font
            title_font.name = 'Malgun Gothic'
            title_font.size = Pt(20)
            title_font.bold = True
            title_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_style.paragraph_format.space_after = Pt(18)
            title_style.paragraph_format.space_before = Pt(12)
        
        # 헤딩 1 스타일 (1. 2. 3... 최상위 번호)
        if 'CustomHeading1' not in styles:
            h1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            h1_font = h1_style.font
            h1_font.name = 'Malgun Gothic'
            h1_font.size = Pt(16)
            h1_font.bold = False
            h1_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            h1_style.paragraph_format.space_before = Pt(18)
            h1_style.paragraph_format.space_after = Pt(12)
            h1_style.paragraph_format.left_indent = Inches(0)
            h1_style.paragraph_format.keep_with_next = True
        
        # 헤딩 2 스타일 (1.1 1.2... 2단계 번호)
        if 'CustomHeading2' not in styles:
            h2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            h2_font = h2_style.font
            h2_font.name = 'Malgun Gothic'
            h2_font.size = Pt(14)
            h2_font.bold = False
            h2_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            h2_style.paragraph_format.space_before = Pt(12)
            h2_style.paragraph_format.space_after = Pt(8)
            h2_style.paragraph_format.left_indent = Inches(0.2)
        
        # 헤딩 3 스타일 (1.1.1 1.1.2... 3단계 번호)
        if 'CustomHeading3' not in styles:
            h3_style = styles.add_style('CustomHeading3', WD_STYLE_TYPE.PARAGRAPH)
            h3_font = h3_style.font
            h3_font.name = 'Malgun Gothic'
            h3_font.size = Pt(12)
            h3_font.bold = False
            h3_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            h3_style.paragraph_format.space_before = Pt(8)
            h3_style.paragraph_format.space_after = Pt(6)
            h3_style.paragraph_format.left_indent = Inches(0.4)
        
        # 본문 스타일 (표준 단락)
        if 'CustomBody' not in styles:
            body_style = styles.add_style('CustomBody', WD_STYLE_TYPE.PARAGRAPH)
            body_font = body_style.font
            body_font.name = 'Malgun Gothic'
            body_font.size = Pt(11)
            body_style.paragraph_format.line_spacing = 1.3
            body_style.paragraph_format.space_after = Pt(6)
            body_style.paragraph_format.first_line_indent = Inches(0.2)
            body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # 리스트 스타일
        if 'CustomList' not in styles:
            list_style = styles.add_style('CustomList', WD_STYLE_TYPE.PARAGRAPH)
            list_font = list_style.font
            list_font.name = 'Malgun Gothic'
            list_font.size = Pt(11)
            list_style.paragraph_format.left_indent = Inches(0.4)
            try:
                list_style.paragraph_format.hanging_indent = Inches(0.2)
            except AttributeError:
                list_style.paragraph_format.left_indent = Inches(0.2)
            list_style.paragraph_format.space_after = Pt(3)
            list_style.paragraph_format.line_spacing = 1.2
        
        # 목차 제목 스타일
        if 'TOCHeading' not in styles:
            toc_heading_style = styles.add_style('TOCHeading', WD_STYLE_TYPE.PARAGRAPH)
            toc_heading_font = toc_heading_style.font
            toc_heading_font.name = 'Malgun Gothic'
            toc_heading_font.size = Pt(16)
            toc_heading_font.bold = True
            toc_heading_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            toc_heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            toc_heading_style.paragraph_format.space_after = Pt(18)
        
        # 목차 항목 스타일 - 레벨 1
        if 'TOCEntry1' not in styles:
            toc_entry1_style = styles.add_style('TOCEntry1', WD_STYLE_TYPE.PARAGRAPH)
            toc_entry1_font = toc_entry1_style.font
            toc_entry1_font.name = 'Malgun Gothic'
            toc_entry1_font.size = Pt(11)
            toc_entry1_style.paragraph_format.space_after = Pt(3)
            toc_entry1_style.paragraph_format.left_indent = Inches(0)
            # 탭 스톱 추가 - 점선으로 페이지 번호까지 연결
            toc_entry1_style.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        
        # 목차 항목 스타일 - 레벨 2
        if 'TOCEntry2' not in styles:
            toc_entry2_style = styles.add_style('TOCEntry2', WD_STYLE_TYPE.PARAGRAPH)
            toc_entry2_font = toc_entry2_style.font
            toc_entry2_font.name = 'Malgun Gothic'
            toc_entry2_font.size = Pt(10)
            toc_entry2_style.paragraph_format.space_after = Pt(2)
            toc_entry2_style.paragraph_format.left_indent = Inches(0.3)
            # 탭 스톱 추가
            toc_entry2_style.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        
        # 목차 항목 스타일 - 레벨 3
        if 'TOCEntry3' not in styles:
            toc_entry3_style = styles.add_style('TOCEntry3', WD_STYLE_TYPE.PARAGRAPH)
            toc_entry3_font = toc_entry3_style.font
            toc_entry3_font.name = 'Malgun Gothic'
            toc_entry3_font.size = Pt(10)
            toc_entry3_style.paragraph_format.space_after = Pt(2)
            toc_entry3_style.paragraph_format.left_indent = Inches(0.6)
            # 탭 스톱 추가
            toc_entry3_style.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
            
        # 일반 목차 스타일 (하위 호환성)
        if 'TOCEntry' not in styles:
            toc_style = styles.add_style('TOCEntry', WD_STYLE_TYPE.PARAGRAPH)
            toc_font = toc_style.font
            toc_font.name = 'Malgun Gothic'
            toc_font.size = Pt(11)
            toc_style.paragraph_format.space_after = Pt(3)
            # 탭 스톱 추가
            toc_style.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        
        # 불릿 리스트 레벨별 스타일
        if 'CustomListLevel1' not in styles:
            level1_style = styles.add_style('CustomListLevel1', WD_STYLE_TYPE.PARAGRAPH)
            level1_font = level1_style.font
            level1_font.name = 'Malgun Gothic'
            level1_font.size = Pt(11)
            level1_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            level1_style.paragraph_format.left_indent = Inches(0.25)
            level1_style.paragraph_format.space_after = Pt(3)
            level1_style.paragraph_format.line_spacing = 1.2
        
        if 'CustomListLevel2' not in styles:
            level2_style = styles.add_style('CustomListLevel2', WD_STYLE_TYPE.PARAGRAPH)
            level2_font = level2_style.font
            level2_font.name = 'Malgun Gothic'
            level2_font.size = Pt(11)
            level2_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            level2_style.paragraph_format.left_indent = Inches(0.5)
            level2_style.paragraph_format.space_after = Pt(3)
            level2_style.paragraph_format.line_spacing = 1.2
        
        if 'CustomListLevel3' not in styles:
            level3_style = styles.add_style('CustomListLevel3', WD_STYLE_TYPE.PARAGRAPH)
            level3_font = level3_style.font
            level3_font.name = 'Malgun Gothic'
            level3_font.size = Pt(11)
            level3_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            level3_style.paragraph_format.left_indent = Inches(0.75)
            level3_style.paragraph_format.space_after = Pt(3)
            level3_style.paragraph_format.line_spacing = 1.2
        
        if 'CustomListLevel4' not in styles:
            level4_style = styles.add_style('CustomListLevel4', WD_STYLE_TYPE.PARAGRAPH)
            level4_font = level4_style.font
            level4_font.name = 'Malgun Gothic'
            level4_font.size = Pt(11)
            level4_font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            level4_style.paragraph_format.left_indent = Inches(1.0)
            level4_style.paragraph_format.space_after = Pt(3)
            level4_style.paragraph_format.line_spacing = 1.2
        
        # 캡션 스타일
        if 'CustomCaption' not in styles:
            caption_style = styles.add_style('CustomCaption', WD_STYLE_TYPE.PARAGRAPH)
            caption_font = caption_style.font
            caption_font.name = 'Malgun Gothic'
            caption_font.size = Pt(10)
            caption_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_style.paragraph_format.space_after = Pt(12)
    
    def add_title(self, title: str):
        """제목 추가"""
        para = self.document.add_paragraph()
        para.style = 'CustomTitle'
        para.add_run(title)
    
    def add_heading(self, text: str, level: int):
        """헤딩 추가"""
        para = self.document.add_paragraph()
        if level == 1:
            para.style = 'CustomHeading1'
        elif level == 2:
            para.style = 'CustomHeading2'
        else:
            para.style = 'CustomHeading3'
        para.add_run(text)
    
    def add_paragraph(self, text: str):
        """단락 추가"""
        # 캡션 처리
        caption_match = re.match(r'<(표|그림)\s*(\d+)>\s*(.*)', text)
        if caption_match:
            caption_type, caption_num, caption_text = caption_match.groups()
            full_caption = f"<{caption_type} {caption_num}> {caption_text}"
            
            if full_caption not in self.processed_captions:
                self.processed_captions.add(full_caption)
                self.add_caption(caption_text, caption_type)
                return
            else:
                return
        
        # 이미지 처리
        img_match = re.search(r'!\[(.*?)\]\((.*?)\)', text)
        if img_match:
            alt_text, image_path = img_match.groups()
            self.add_image(image_path, alt_text)
            return
        
        # 일반 텍스트 처리
        if text.strip():
            para = self.document.add_paragraph()
            para.style = 'CustomBody'
            
            # 볼드/이탤릭 처리
            self.process_inline_formatting(text, para)
    
    def _has_footnote_pattern(self, text: str) -> bool:
        """주석 패턴이 있는지 확인 (^숫자^[내용] 또는 ^숫자^ 형태)"""
        return bool(re.search(r'\^\d+\^(\[.*?\])?', text))
    
    def _add_paragraph_with_footnotes(self, text: str, style: str = 'CustomBody'):
        """주석이 포함된 단락 추가"""
        para = self.document.add_paragraph()
        para.style = style
        
        # 주석 패턴 분할: ^숫자^[내용] 또는 ^숫자^ 형태
        footnote_pattern = r'\^(\d+)\^(\[([^\]]*)\])?'
        
        last_end = 0
        
        for match in re.finditer(footnote_pattern, text):
            start, end = match.span()
            
            # 매치 이전 텍스트 추가
            if start > last_end:
                para.add_run(text[last_end:start])
            
            # 상첨자 번호 추가
            footnote_num = match.group(1)
            run = para.add_run(footnote_num)
            run.font.superscript = True
            
            # 대괄호 안의 내용이 있으면 추가 (각주 내용)
            footnote_content = match.group(3)
            if footnote_content:
                footnote_run = para.add_run(f"[{footnote_content}]")
                footnote_run.font.size = Pt(9)  # 각주는 작은 크기
            
            last_end = end
        
        # 나머지 텍스트 추가
        if last_end < len(text):
            para.add_run(text[last_end:])
        
        return para

    def process_inline_formatting(self, text: str, para):
        """인라인 포맷팅 처리 - 상첨자 지원 추가"""
        # 상첨자가 있는 경우 특별 처리
        if self._has_footnote_pattern(text):
            # 기존 단락 제거하고 새로 만들기
            if para in self.document.paragraphs:
                p = para._element
                p.getparent().remove(p)
            return self._add_paragraph_with_footnotes(text, para.style.name if hasattr(para, 'style') else 'CustomBody')
        
        # 간단한 볼드 처리
        parts = re.split(r'(\*\*.*?\*\*)', text)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = para.add_run(part[2:-2])
                run.bold = True
            else:
                para.add_run(part)
    
    def add_caption(self, caption_text: str, caption_type: str = "Figure"):
        """캡션 추가"""
        if caption_type.lower() == "figure" or caption_type.lower() == "그림":
            self.figure_counter += 1
            caption_num = self.figure_counter
            label = "그림"
        else:
            self.table_counter += 1
            caption_num = self.table_counter
            label = "표"
        
        caption_para = self.document.add_paragraph()
        caption_para.style = 'CustomCaption'
        full_caption = f"<{label} {caption_num}> {caption_text}"
        caption_para.add_run(full_caption)
    
    def add_image(self, image_path: str, alt_text: str = ""):
        """이미지 추가"""
        if os.path.exists(image_path):
            try:
                para = self.document.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.runs[0] if para.runs else para.add_run()
                run.add_picture(image_path, width=Inches(5))
            except Exception as e:
                # 이미지 추가 실패시 텍스트로 대체
                para = self.document.add_paragraph()
                para.style = 'CustomBody'
                para.add_run(f"[이미지: {alt_text}]")
    
    def add_table(self, headers: List[str], rows: List[List[str]], caption_text: str = None):
        """테이블 추가 - 캡션 통합 지원"""
        if not headers or not rows:
            return
        
        table = self.document.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = 'Table Grid'
        
        # 헤더 설정
        for i, header in enumerate(headers):
            cell = table.cell(0, i)
            cell.text = header
            # 헤더 스타일
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = 'Malgun Gothic'
                    run.font.size = Pt(10)
        
        # 데이터 행 설정
        for row_idx, row_data in enumerate(rows):
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < len(headers):
                    cell = table.cell(row_idx + 1, col_idx)
                    cell.text = str(cell_data)
                    # 셀 스타일
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.name = 'Malgun Gothic'
                            run.font.size = Pt(10)
        
        # 테이블 생성 후 캡션 추가 (아래쪽에)
        if caption_text:
            self.add_caption(caption_text, "표")
    
    def parse_table_from_lines(self, table_lines):
        """마크다운 테이블 라인들을 파싱하여 테이블 추가"""
        if len(table_lines) < 2:
            return
        
        # 헤더 파싱
        header_line = table_lines[0].strip()
        headers = [cell.strip() for cell in header_line.split('|')[1:-1]]
        
        # 데이터 행 파싱
        rows = []
        for line in table_lines[2:]:  # 헤더와 구분선 다음부터
            if line.strip():
                cells = [cell.strip() for cell in line.strip().split('|')[1:-1]]
                rows.append(cells)
        
        if rows:
            self.add_table(headers, rows)
    
    def add_list(self, items: List[str], ordered: bool = False):
        """리스트 추가"""
        for item in items:
            para = self.document.add_paragraph()
            para.style = 'CustomList'
            
            if ordered:
                # 번호가 있는 리스트는 그대로
                para.add_run(item)
            else:
                # 기호 추가
                para.add_run(f"• {item}")
    
    def process_toc_line(self, line: str):
        """목차 라인 처리 - 완전한 점선과 페이지 번호 지원"""
        # 목차 항목 식별 및 처리 - 다양한 형태 지원
        # 형태 1: "　　1.1 소제목 .......... 3"
        # 형태 2: "1. 제목……………………5"
        # 형태 3: "1.1.1 세부제목 ... 10"
        
        toc_pattern = r'^(　*)([\d\.]+\s+.*?)\s*[…\.]*\s*(\d*)\s*$'
        match = re.match(toc_pattern, line)
        
        if match:
            indent_spaces, content, page_num = match.groups()
            
            # 레벨 결정 (점의 개수로 판단)
            level = 1
            if '.' in content:
                dots = content.split()[0].count('.')
                level = min(dots + 1, 3)  # 최대 3레벨
            
            # 들여쓰기 개수로도 레벨 보정
            indent_count = len(indent_spaces)
            if indent_count > 0:
                level = min(indent_count + 1, 3)
            
            para = self.document.add_paragraph()
            
            # 레벨별 스타일 적용
            if level == 1:
                para.style = 'TOCEntry1'
            elif level == 2:
                para.style = 'TOCEntry2'
            else:
                para.style = 'TOCEntry3'
            
            # 콘텐츠 추가 (제목 부분)
            title_text = content.strip()
            para.add_run(title_text)
            
            # 탭 문자 추가 (점선 리더와 함께)
            para.add_run('\t')
            
            # 페이지 번호 추가 (있는 경우)
            if page_num:
                para.add_run(page_num)
            else:
                para.add_run('1')  # 기본 페이지 번호
            
            return True
        return False
    
    def convert_markdown(self, md_content: str):
        """원본 마크다운을 직접 라인별로 처리 - 원래 로직 기반"""
        lines = md_content.split('\n')
        i = 0
        current_table = []
        in_table = False
        current_table_caption = None
        
        while i < len(lines):
            original_line = lines[i] 
            line = lines[i].strip()
            
            # 빈 줄 건너뛰기
            if not line:
                i += 1
                continue
            
            # 구분선 건너뛰기
            if line == '---':
                i += 1
                continue
            
            # 이미지 처리
            if self._is_image_line(line):
                self._add_image(line)
                i += 1
                continue
            
            # 테이블 처리
            if '|' in line and not in_table:
                # 테이블 시작 전에 이전 라인이 표 캡션인지 확인
                table_caption_before = None
                table_caption_before_text = None
                if i > 0:
                    prev_line = lines[i - 1].strip()
                    if re.match(r'^<표\s*\d+>\s*.*', prev_line, re.IGNORECASE):
                        table_caption_before = prev_line
                        
                        # 캡션이 아직 처리되지 않았으면 텍스트 추출
                        if table_caption_before not in self.processed_captions:
                            match = re.match(r'^<표\s*\d+>\s*(.*)', table_caption_before, re.IGNORECASE)
                            if match:
                                table_caption_before_text = match.group(1).strip()
                                self.processed_captions.add(table_caption_before)
                
                # 테이블 시작
                in_table = True
                current_table = [line]
                current_table_caption = table_caption_before_text  # 이전에 발견된 캡션 저장
                i += 1
                continue
            elif in_table:
                if '|' in line:
                    current_table.append(line)
                    i += 1
                    continue
                else:
                    # 테이블 끝 - 다음 최대 3줄에서 표 캡션 찾기
                    table_caption = None
                    caption_line_index = None
                    
                    for j in range(1, 4):  # 최대 3줄까지 확인
                        if i + j < len(lines):
                            check_line = lines[i + j].strip()
                            if not check_line:  # 빈 라인은 건너뛰기
                                continue
                            if re.match(r'^<표\s*\d+>\s*.*', check_line, re.IGNORECASE):
                                table_caption = check_line
                                caption_line_index = i + j
                                break
                            else:
                                # 다른 텍스트가 나오면 캡션 찾기 중단
                                break
                    
                    # 테이블 캡션 텍스트 추출
                    caption_text = None
                    if table_caption:
                        match = re.match(r'^<표\s*\d+>\s*(.*)', table_caption, re.IGNORECASE)
                        if match:
                            caption_text = match.group(1).strip()
                            self.processed_captions.add(table_caption)
                            
                            # 캡션 라인을 건너뛰도록 인덱스 조정
                            if caption_line_index:
                                i = caption_line_index
                    
                    # 최종 캡션 결정 - 테이블 후에 찾은 캡션이 우선, 없으면 이전 캡션 사용
                    final_caption = caption_text if caption_text else current_table_caption
                    
                    # 테이블 생성 (캡션과 함께)
                    self._create_table_from_lines(current_table, final_caption)
                    current_table = []
                    current_table_caption = None
                    in_table = False
                    
            # 헤딩 처리 (# ## ###)
            elif line.startswith('#'):
                self._process_heading_line(line)
            # 목차 항목 처리 (점선과 페이지 번호가 있는 경우) - 우선 처리
            elif self._is_toc_entry(original_line):
                self._add_toc_entry(original_line)
            # 계층적 번호 체계 처리
            elif self._is_hierarchical_content(line):
                self._add_hierarchical_content(line)
            # 로마숫자로 시작하는 대제목 (Ⅰ, Ⅱ, Ⅲ...)
            elif re.match(r'^[ⅠⅡⅢⅣⅤⅥⅦⅧⅨⅩⅪⅫⅩⅢⅩⅣⅩⅤⅩⅥⅩⅦⅩⅧⅩⅨⅩⅩ]+\s+', line):
                para = self.document.add_paragraph(line)
                para.style = 'CustomHeading1'
            # 괄호 숫자로 시작하는 소제목 (1) 2) 3))
            elif re.match(r'^\d+\)\s+', line):
                para = self.document.add_paragraph(line)
                para.style = 'CustomHeading3'
            # 불릿 포인트 처리
            elif line.startswith('□'):
                para = self.document.add_paragraph(line)
                para.style = 'CustomListLevel1'
            elif line.startswith('○'):
                para = self.document.add_paragraph(line)
                para.style = 'CustomListLevel2'  
            elif line.startswith('-') and not line.startswith('---'):
                para = self.document.add_paragraph(line)
                para.style = 'CustomListLevel3'
            elif line.startswith('•'):
                para = self.document.add_paragraph(line)
                para.style = 'CustomListLevel4'
            # 특수 번호 (①②③)
            elif re.match(r'^[①②③④⑤⑥⑦⑧⑨⑩]', line):
                para = self.document.add_paragraph(line)
                para.style = 'CustomList'
            # 캡션 처리
            elif re.match(r'^<(표|그림)\s*\d+>\s*.*', line):
                caption_match = re.match(r'^<(표|그림)\s*\d+>\s*(.*)', line)
                if caption_match:
                    caption_type, caption_text = caption_match.groups()
                    full_caption = f"<{caption_type} {self.figure_counter+1 if caption_type == '그림' else self.table_counter+1}> {caption_text}"
                    
                    if full_caption not in self.processed_captions:
                        self.processed_captions.add(full_caption)
                        self.add_caption(caption_text, caption_type)
            # 일반 텍스트
            else:
                # 상첨자가 있는 텍스트는 특별 처리
                if self._has_footnote_pattern(line):
                    self._add_paragraph_with_footnotes(line, 'CustomBody')
                else:
                    para = self.document.add_paragraph(line)
                    para.style = 'CustomBody'
            
            i += 1
        
        # 마지막에 테이블이 있으면 처리
        if current_table:
            self._create_table_from_lines(current_table, current_table_caption)
    
    def _is_image_line(self, line: str) -> bool:
        """이미지 라인인지 확인"""
        return re.match(r'!\[.*?\]\(.*?\)', line) is not None
    
    def _add_image(self, line: str):
        """이미지 추가"""
        match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if match:
            alt_text, image_path = match.groups()
            self.add_image(image_path, alt_text)
    
    def _process_heading_line(self, line: str):
        """헤딩 라인 처리"""
        if line.startswith('###'):
            text = line[3:].strip()
            para = self.document.add_paragraph(text)
            para.style = 'CustomHeading3'
        elif line.startswith('##'):
            text = line[2:].strip()
            para = self.document.add_paragraph(text)
            para.style = 'CustomHeading2'
        elif line.startswith('#'):
            text = line[1:].strip()
            para = self.document.add_paragraph(text)
            para.style = 'CustomTitle'
    
    def _is_toc_entry(self, line: str) -> bool:
        """목차 항목인지 확인 - 향상된 패턴 매칭"""
        # 다양한 목차 패턴 지원
        # 1. "1. 제목......5" 형태
        # 2. "　1.1 소제목…3" 형태  
        # 3. "1.1.1 세부제목 ... 10" 형태
        toc_pattern = r'^(　*)([\d\.]+\s+.*?)\s*[…\.]*\s*(\d*)\s*$'
        
        # 추가 패턴: 번호 + 제목 + 점선/생략표 + 페이지번호
        advanced_pattern = r'^(　*\d+[\.\d]*\s+.*?)\s*[…\.]{2,}\s*\d+\s*$'
        
        return (re.match(toc_pattern, line) is not None or 
                re.match(advanced_pattern, line) is not None)
    
    def _add_toc_entry(self, line: str):
        """목차 항목 추가"""
        self.process_toc_line(line)
    
    def _is_hierarchical_content(self, line: str) -> bool:
        """계층적 번호 체계 콘텐츠인지 확인"""
        # 1. 형태의 번호 (1단계)
        if re.match(r'^\d+\.\s+', line):
            return True
        # 1.1 형태의 번호 (2단계)
        if re.match(r'^\d+\.\d+\s+', line):
            return True
        # 1.1.1 형태의 번호 (3단계)
        if re.match(r'^\d+\.\d+\.\d+\s+', line):
            return True
        return False
    
    def _add_hierarchical_content(self, line: str):
        """계층적 번호 체계 콘텐츠 추가 - 상첨자 지원"""
        # 상첨자가 있는 경우 특별 처리
        if self._has_footnote_pattern(line):
            if re.match(r'^\d+\.\s+', line):
                self._add_paragraph_with_footnotes(line, 'CustomHeading1')
            elif re.match(r'^\d+\.\d+\s+', line):
                self._add_paragraph_with_footnotes(line, 'CustomHeading2')
            elif re.match(r'^\d+\.\d+\.\d+\s+', line):
                self._add_paragraph_with_footnotes(line, 'CustomHeading3')
        else:
            para = self.document.add_paragraph(line)
            
            if re.match(r'^\d+\.\s+', line):
                para.style = 'CustomHeading1'
            elif re.match(r'^\d+\.\d+\s+', line):
                para.style = 'CustomHeading2'
            elif re.match(r'^\d+\.\d+\.\d+\s+', line):
                para.style = 'CustomHeading3'
    
    def _add_toc_heading(self, title: str = "목   차"):
        """목차 제목 추가"""
        para = self.document.add_paragraph()
        para.style = 'TOCHeading'
        para.add_run(title)
        return para
    
    def _insert_page_break(self):
        """페이지 브레이크 삽입"""
        para = self.document.add_paragraph()
        run = para.runs[0] if para.runs else para.add_run()
        run.add_break(WD_BREAK.PAGE)
        return para
        
    def _add_cover_title(self, title: str):
        """표지 제목 추가"""
        para = self.document.add_paragraph()
        para.style = 'CoverTitle'
        para.add_run(title)
        return para
    
    def _create_table_from_lines(self, table_lines, caption_text: str = None):
        """테이블 라인들로부터 테이블 생성 - 캡션 지원"""
        if len(table_lines) < 2:
            return
        
        # 헤더 파싱
        header_line = table_lines[0].strip()
        headers = [cell.strip() for cell in header_line.split('|')[1:-1]]
        
        # 데이터 행 파싱 (구분선 건너뛰기)
        rows = []
        for line in table_lines[2:]:  # 헤더와 구분선 다음부터
            if line.strip() and '|' in line:
                cells = [cell.strip() for cell in line.strip().split('|')[1:-1]]
                rows.append(cells)
        
        if rows:
            self.add_table(headers, rows, caption_text)
    
    def save(self, output_path: str):
        """문서 저장"""
        self.document.save(output_path)


def main():
    if len(sys.argv) < 2:
        print("사용법: python simple_converter.py input.md [output.docx]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.md', '.docx')
    
    if not os.path.exists(input_file):
        print(f"입력 파일을 찾을 수 없습니다: {input_file}")
        sys.exit(1)
    
    try:
        converter = SimpleDocxConverter()
        
        with open(input_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        converter.convert_markdown(content)
        converter.save(output_file)
        
        print(f"변환 완료: {output_file}")
    
    except Exception as e:
        print(f"변환 중 오류가 발생했습니다: {str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()