#!/usr/bin/env python3
"""
Universal MD to DOCX Converter
모든 사업계획서와 문서에 범용적으로 사용 가능한 변환기
"""

import os
import re
from typing import List, Optional
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

class UniversalMDConverter:
    def __init__(self):
        self.document = Document()
        self.setup_styles()
        
    def setup_styles(self):
        """Word 스타일 설정"""
        styles = self.document.styles
        
        # 제목 스타일
        if 'CustomHeading1' not in styles:
            h1_style = styles.add_style('CustomHeading1', WD_STYLE_TYPE.PARAGRAPH)
            h1_font = h1_style.font
            h1_font.name = 'Arial'
            h1_font.size = Pt(16)
            h1_font.bold = True
            h1_style.paragraph_format.space_before = Pt(18)
            h1_style.paragraph_format.space_after = Pt(12)
            
        if 'CustomHeading2' not in styles:
            h2_style = styles.add_style('CustomHeading2', WD_STYLE_TYPE.PARAGRAPH)
            h2_font = h2_style.font  
            h2_font.name = 'Arial'
            h2_font.size = Pt(14)
            h2_font.bold = True
            h2_style.paragraph_format.space_before = Pt(12)
            h2_style.paragraph_format.space_after = Pt(8)
            
    def convert(self, md_file: str) -> str:
        """메인 변환 함수 - Windows 순서: 텍스트 먼저, 이미지는 나중에"""
        print(f"🔄 변환 시작: {md_file}")
        
        # MD 파일의 디렉토리 저장 (이미지 경로 처리용)
        self.md_file_dir = os.path.dirname(os.path.abspath(md_file))
        print(f"🔍 MD 파일 경로 디버깅:")
        print(f"   원본 경로: {md_file}")
        print(f"   절대 경로: {os.path.abspath(md_file)}")
        print(f"   디렉토리: {self.md_file_dir}")
        print(f"   디렉토리 존재: {os.path.exists(self.md_file_dir)}")
        
        with open(md_file, 'r', encoding='utf-8') as f:
            content = f.read()
            
        lines = content.split('\n')
        
        # 이미지 정보를 저장할 리스트 (나중에 삽입용)
        self.pending_images = []
        
        # MD 파일의 첫 번째 # 제목 찾기
        main_title = None
        for line in lines:
            line_stripped = line.strip()
            if line_stripped.startswith('# '):
                main_title = line_stripped[2:].strip()
                break
        
        # 제목 추가 (MD에서 찾은 제목 또는 기본값)
        if main_title:
            title_para = self.document.add_paragraph(main_title)
        else:
            title_para = self.document.add_paragraph("문서 제목")
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.name = 'Arial'
        title_run.font.size = Pt(20)
        title_run.font.bold = True
        
        self.document.add_page_break()
        
        i = 0
        while i < len(lines):
            line = lines[i]  # 원래 라인 (들여쓰기 포함)
            line_stripped = line.strip()
            
            if not line_stripped:  # 빈 줄
                i += 1
                continue
                
            if line_stripped.startswith('# '):  # H1 제목 (문서 제목이므로 스킵)
                i += 1
                continue
                
            elif line_stripped.startswith('## '):  # H2 제목  
                title = line_stripped[3:].strip()
                if title == '주석':  # 주석 섹션은 따로 처리
                    i = self.process_footnote_section(lines, i)
                    continue
                else:
                    para = self.document.add_paragraph(title)
                    para.style = 'CustomHeading1'
                    
            elif line_stripped.startswith('### '):  # H3 제목
                title = line_stripped[4:].strip()
                para = self.document.add_paragraph(title)
                para.style = 'CustomHeading2'
                
            elif line_stripped.startswith('#### '):  # H4 제목
                title = line_stripped[5:].strip()
                para = self.document.add_paragraph(title)
                run = para.runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(12)
                run.font.bold = True
                para.paragraph_format.space_before = Pt(8)
                para.paragraph_format.space_after = Pt(4)
                
            elif line_stripped.startswith('!['):  # 이미지 - 플레이스홀더 삽입 후 나중에 교체
                i = self.store_image_with_placeholder(lines, i)
                continue
                
            elif line_stripped.startswith('<그림') or line_stripped.startswith('<표'):  # 캡션
                para = self.document.add_paragraph(line_stripped)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.bold = True
                
            elif line_stripped.startswith('|') and '|' in line_stripped:  # 테이블
                i = self.process_table(lines, i)
                continue
                
            elif line_stripped.startswith('□') or line_stripped.startswith('○') or line_stripped.startswith('-') or line_stripped.startswith('•'):  # 불릿 포인트
                # 들여쓰기 레벨 계산
                indent_level = self.get_bullet_level(line)
                self.add_bullet_paragraph(line, indent_level)  # 원래 line 전달 (들여쓰기 포함)
                
            elif line_stripped != '---':  # 일반 텍스트 (구분선 제외)
                if line_stripped:  # 빈 줄이 아닌 경우만
                    # 주석 참조 처리 (^1^, ^2^ 등)
                    processed_line = self.process_footnote_references(line)
                    para = self.document.add_paragraph(processed_line)
                    run = para.runs[0]
                    run.font.name = 'Arial'
                    run.font.size = Pt(11)
                    
            i += 1
            
        # 1단계: 텍스트만으로 DOCX 저장
        import time
        timestamp = int(time.time())
        output_filename = os.path.basename(md_file).replace('.md', f'_TEST_{timestamp}.docx')
        
        print(f"🔍 저장 경로 디버깅:")
        print(f"   MD 파일: {md_file}")
        print(f"   MD 디렉토리: {self.md_file_dir}")
        print(f"   출력 파일명: {output_filename}")
        
        # MD 파일과 같은 디렉토리에 저장
        output_file = os.path.join(self.md_file_dir, output_filename)
        print(f"   최종 저장 경로: {output_file}")
        
        self.document.save(output_file)
        print(f"✅ 1단계 완료: 텍스트 DOCX 저장")
        
        # 2단계: 이미지들 삽입
        if self.pending_images:
            print(f"🖼️  2단계 시작: {len(self.pending_images)}개 이미지 삽입")
            self.insert_pending_images()
            self.document.save(output_file)  # 이미지 삽입 후 재저장
            print(f"✅ 2단계 완료: 이미지 삽입 후 재저장")
        
        # 3단계: 주석 섹션 자동 생성
        self.add_footnotes_from_content()
        self.document.save(output_file)  # 주석 추가 후 재저장
        print(f"✅ 3단계 완료: 주석 섹션 추가")
        
        print(f"✅ 변환 완료: {output_file}")
        return output_file
        
    def process_footnote_section(self, lines: List[str], start_idx: int) -> int:
        """주석 섹션 처리 - 중복 방지"""
        print("📝 주석 섹션 처리 중...")
        
        # 주석 제목 추가 (한 번만!)
        self.document.add_page_break()
        title_para = self.document.add_paragraph("주석")
        title_para.style = 'CustomHeading1'
        
        i = start_idx + 1  # "## 주석" 다음 라인부터
        while i < len(lines):
            line = lines[i].strip()
            
            if not line:
                i += 1
                continue
                
            if line.startswith('#'):  # 다음 섹션 시작
                break
                
            # 주석 내용 추가
            para = self.document.add_paragraph(line)
            run = para.runs[0]
            run.font.name = 'Arial'  
            run.font.size = Pt(11)  # 동일한 크기로 통일
            
            i += 1
            
        return i
    
    def store_image_with_placeholder(self, lines: List[str], start_idx: int) -> int:
        """이미지 위치에 플레이스홀더를 삽입하고 나중에 교체"""
        line = lines[start_idx].strip()
        
        # ![alt](path) 형식 파싱
        match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if match:
            alt_text = match.group(1)
            image_path = match.group(2)
            
            # 다음 줄 캡션 확인
            next_caption = None
            next_idx = start_idx + 1
            if next_idx < len(lines):
                next_line = lines[next_idx].strip()
                if next_line.startswith('<그림'):
                    next_caption = next_line
            
            # 플레이스홀더 문단 생성 (나중에 이미지로 교체)
            placeholder_id = f"IMAGE_PLACEHOLDER_{len(self.pending_images)}"
            placeholder_para = self.document.add_paragraph(f"[{placeholder_id}]")
            placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 이미지 정보 저장 (플레이스홀더 문단 참조 포함)
            image_info = {
                'alt_text': alt_text,
                'image_path': image_path,
                'next_caption': next_caption,
                'placeholder_para': placeholder_para,  # 플레이스홀더 문단 참조
                'placeholder_id': placeholder_id
            }
            self.pending_images.append(image_info)
            print(f"📋 이미지 플레이스홀더 삽입: {image_path} -> {placeholder_id}")
            
            # 캡션이 다음 줄에 있으면 건너뛰기
            if next_caption:
                return next_idx + 1
        
        return start_idx + 1
    
    def process_footnote_references(self, line: str) -> str:
        """주석 참조 ^1^[설명] 형태를 처리"""
        import re
        
        # ^숫자^[설명] 패턴을 찾아서 상위첨자로 변환
        def replace_footnote(match):
            number = match.group(1)
            description = match.group(2)
            
            # 주석 번호만 반환 (설명은 제거)
            return f"({number})"
        
        # ^숫자^[설명] 패턴 처리
        processed = re.sub(r'\^(\d+)\^\[([^\]]+)\]', replace_footnote, line)
        
        return processed
    
    def add_footnotes_from_content(self):
        """MD 내용에서 주석을 추출하여 주석 섹션 생성"""
        # 현재 처리 중인 MD 파일에서 주석 정의 찾기
        footnotes = {}
        
        # 현재 MD 파일 경로 동적 확인
        md_files = [f for f in os.listdir(self.md_file_dir) if f.endswith('.md')]
        if not md_files:
            print("⚠️  MD 파일을 찾을 수 없음")
            return
            
        md_file_path = os.path.join(self.md_file_dir, md_files[0])  # 첫 번째 MD 파일 사용
        
        with open(md_file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            
        import re
        # ^숫자^[설명] 패턴 찾기
        footnote_pattern = r'\^(\d+)\^\[([^\]]+)\]'
        matches = re.findall(footnote_pattern, content)
        
        for number, description in matches:
            if number not in footnotes:  # 중복 방지
                footnotes[number] = description
        
        if footnotes:
            print(f"📝 주석 {len(footnotes)}개 발견, 주석 섹션 생성 중...")
            
            # 주석 섹션 추가
            self.document.add_page_break()
            title_para = self.document.add_paragraph("주석")
            title_para.style = 'CustomHeading1'
            
            # 번호 순으로 정렬하여 주석 추가
            for number in sorted(footnotes.keys(), key=int):
                description = footnotes[number]
                footnote_text = f"({number}) {description}"
                
                para = self.document.add_paragraph(footnote_text)
                run = para.runs[0]
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)
                
            print(f"✅ 주석 섹션 완료: {len(footnotes)}개 주석 추가")
        else:
            print("📝 주석이 발견되지 않음")
    
    def insert_pending_images(self):
        """플레이스홀더를 실제 이미지로 교체"""
        for img_info in self.pending_images:
            alt_text = img_info['alt_text']
            image_path = img_info['image_path']
            placeholder_para = img_info['placeholder_para']
            placeholder_id = img_info['placeholder_id']
            
            # 절대 경로로 변환
            if not os.path.isabs(image_path):
                full_path = os.path.join(self.md_file_dir, image_path)
            else:
                full_path = image_path
            
            print(f"🖼️  플레이스홀더 교체: {placeholder_id} -> {image_path}")
            
            # 플레이스홀더 문단의 텍스트를 지우고 이미지로 교체
            placeholder_para.clear()
            
            if os.path.exists(full_path):
                try:
                    run = placeholder_para.add_run()
                    run.add_picture(full_path, width=Inches(5))
                    placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    print(f"✅ 이미지 교체 성공: {image_path}")
                except Exception as e:
                    print(f"❌ 이미지 교체 실패: {e}")
                    placeholder_para.add_run(f"[이미지 오류: {alt_text}]")
                    placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                # 캡션 처리 (이미지 바로 다음에 추가)
                if img_info['next_caption']:
                    # 플레이스홀더 다음에 캡션 문단 삽입
                    p = placeholder_para._element
                    caption_p = self.document.add_paragraph(img_info['next_caption'])._element
                    p.addnext(caption_p)
                    
                    # 캡션 스타일 적용
                    caption_para = self.document.paragraphs[-1]  # 방금 추가된 문단
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.runs[0]
                    caption_run.font.name = 'Arial'
                    caption_run.font.size = Pt(10)
                    caption_run.font.bold = True
                    print(f"📝 캡션 추가: {img_info['next_caption']}")
            else:
                print(f"⚠️  이미지 파일 없음: {full_path}")
                placeholder_para.add_run(f"[이미지 없음: {alt_text} - {image_path}]")
                placeholder_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
    def process_image(self, lines: List[str], start_idx: int) -> int:
        """이미지 처리 - MD 파일의 캡션 위치를 그대로 존중"""
        line = lines[start_idx].strip()
        
        # ![alt](path) 형식 파싱
        match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
        if match:
            alt_text = match.group(1)
            image_path = match.group(2)
            
            # 절대 경로로 변환 - MD 파일의 디렉토리를 기준으로
            if not os.path.isabs(image_path):
                full_path = os.path.join(self.md_file_dir, image_path)
            else:
                full_path = image_path
                
            print(f"🖼️  이미지 처리: {image_path} -> {full_path}")
            
            # 이전 줄이 캡션인지 확인 (표 캡션이 위에 있는 경우)
            prev_caption = None
            if start_idx > 0:
                prev_line = lines[start_idx - 1].strip()
                if prev_line.startswith('<표'):
                    prev_caption = prev_line
                    print(f"📝 이전 줄 표 캡션 감지: {prev_caption}")
            
            if os.path.exists(full_path):
                # 이미지 추가
                para = self.document.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                try:
                    run.add_picture(full_path, width=Inches(5))
                    print(f"✅ 이미지 추가 성공: {image_path}")
                except Exception as e:
                    print(f"❌ 이미지 추가 실패: {e}")
                    # 실패시 텍스트로 표시
                    para = self.document.add_paragraph(f"[이미지: {alt_text}]")
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                print(f"⚠️  이미지 파일 없음: {full_path}")
                # 파일이 없으면 텍스트로 표시
                para = self.document.add_paragraph(f"[이미지 없음: {alt_text} - {image_path}]")
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 다음 줄이 그림 캡션인지 확인 (그림 캡션이 아래에 있는 경우)
            next_idx = start_idx + 1
            if next_idx < len(lines) and not prev_caption:  # 이전에 캡션이 없었을 때만
                next_line = lines[next_idx].strip()
                if next_line.startswith('<그림'):
                    # 그림 캡션 추가
                    caption_para = self.document.add_paragraph(next_line)
                    caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    caption_run = caption_para.runs[0]
                    caption_run.font.name = 'Arial'
                    caption_run.font.size = Pt(10)
                    caption_run.font.bold = True
                    print(f"📝 그림 캡션 추가: {next_line}")
                    return next_idx + 1  # 캡션까지 처리했으므로 +2
                
        return start_idx + 1
        
    def process_table(self, lines: List[str], start_idx: int) -> int:
        """테이블 처리"""
        table_lines = []
        i = start_idx
        
        # 테이블 라인들 수집
        while i < len(lines) and lines[i].strip().startswith('|'):
            line = lines[i].strip()
            if not line.startswith('|---'):  # 구분선 제외
                table_lines.append(line)
            i += 1
            
        if len(table_lines) < 2:  # 최소 헤더 + 1행
            return i
            
        # 첫 번째 행에서 열 수 계산
        header_cells = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]
        col_count = len(header_cells)
        
        if col_count == 0:
            return i
            
        # 테이블 생성
        table = self.document.add_table(rows=1, cols=col_count)
        table.style = 'Table Grid'
        
        # 헤더 추가
        header_row = table.rows[0]
        for j, cell_text in enumerate(header_cells):
            if j < len(header_row.cells):
                header_row.cells[j].text = cell_text
                # 헤더 셀 볼드 처리
                for run in header_row.cells[j].paragraphs[0].runs:
                    run.font.bold = True
                    run.font.name = 'Arial'
                    run.font.size = Pt(10)
                
                # 헤더 셀 배경색 설정 (연한 회색)
                from docx.oxml.shared import qn
                from docx.oxml import parse_xml
                cell = header_row.cells[j]
                cell_properties = cell._tc.get_or_add_tcPr()
                shade_element = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(
                    'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'))
                cell_properties.append(shade_element)
                    
        # 데이터 행들 추가  
        for line in table_lines[1:]:
            data_cells = [cell.strip() for cell in line.split('|')[1:-1]]
            if len(data_cells) >= col_count:
                row = table.add_row()
                for j, cell_text in enumerate(data_cells[:col_count]):
                    row.cells[j].text = cell_text
                    # **텍스트** 볼드 처리
                    if '**' in cell_text:
                        cell_text_processed = cell_text.replace('**', '')
                        row.cells[j].text = cell_text_processed
                        for run in row.cells[j].paragraphs[0].runs:
                            run.font.bold = True
                    
                    # 폰트 설정
                    for run in row.cells[j].paragraphs[0].runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(10)
                        
        return i
        
    def get_bullet_level(self, line: str) -> int:
        """불릿포인트의 들여쓰기 레벨 계산"""
        # 앞쪽 공백 개수로 레벨 판단
        stripped = line.lstrip()
        spaces = len(line) - len(stripped)
        
        
        if spaces == 0:  # 들여쓰기 없음 (□)
            return 1
        elif spaces == 2:  # 2칸 들여쓰기 (○)
            return 2  
        elif spaces == 4:  # 4칸 들여쓰기 (-)
            return 3
        elif spaces >= 6:  # 6칸 이상 들여쓰기 (•)
            return 4
        else:
            return 1  # 기본값
    
    def add_bullet_paragraph(self, text: str, level: int = 1):
        """MD 파일의 원래 불릿 기호를 그대로 유지하여 Word에 추가"""
        # 기호 제거하지 않고 MD 텍스트 그대로 사용
        para = self.document.add_paragraph()
        
        # 계층별 들여쓰기만 설정 (기호는 변경하지 않음)
        if level == 1:  # □ 기호 (들여쓰기 없음)
            para.paragraph_format.left_indent = Inches(0)
        elif level == 2:  # ○ 기호 (2칸 들여쓰기)
            para.paragraph_format.left_indent = Inches(0.1)
        elif level == 3:  # - 기호 (4칸 들여쓰기)
            para.paragraph_format.left_indent = Inches(0.2)
        else:  # • 기호 (6칸 이상 들여쓰기)
            para.paragraph_format.left_indent = Inches(0.3)
        
        # MD 파일의 원래 텍스트 그대로 추가 (□, ○, -, • 기호 유지)
        para.add_run(text)
        
        # 폰트 설정
        if para.runs:
            run = para.runs[0]
            run.font.name = 'Arial'
            run.font.size = Pt(11)


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("사용법: python3 universal_md_converter.py <MD파일명>")
        sys.exit(1)
    
    md_file = sys.argv[1]
    if not os.path.exists(md_file):
        print(f"❌ 파일을 찾을 수 없습니다: {md_file}")
        sys.exit(1)
    
    converter = UniversalMDConverter()
    result = converter.convert(md_file)
    print(f"\n🎉 변환 완료!\n파일: {result}")